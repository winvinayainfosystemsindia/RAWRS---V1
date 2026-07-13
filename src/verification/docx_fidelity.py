"""DOCX Fidelity (M-3.3) — the simplest useful comparison of a generated
DOCX against its expected benchmark counterpart.

Deliberately NOT a visual/rendering diff (out of scope per the M-3.3
plan) — a structural counts comparison only, using python-docx (already
a project dependency — see src/docx/docx_generator.py), the same library
that writes these files in the first place.

Fidelity formula (kept deterministic and simple, documented here since
that's the one place it can't drift from the code):

    fidelity = 1 - (sum of abs(generated[k] - expected[k]) for every
                    count category k) / max(1, sum of expected[k])

1.0 means every category matches exactly; each mismatched unit in any
category subtracts proportionally to the expected document's total
structural element count. Clamped to [0.0, 1.0].
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Union

from docx import Document as DocxDocument
from docx.oxml.ns import qn

_COUNT_CATEGORIES = ("heading_count", "paragraph_count", "table_count", "figure_count", "page_break_count")


def _counts(docx_path: Union[str, Path]) -> Dict[str, int]:
    doc = DocxDocument(str(docx_path))
    heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
    page_break_count = sum(
        1
        for p in doc.paragraphs
        for r in p.runs
        for br in r._element.findall(qn("w:br"))
        if br.get(qn("w:type")) == "page"
    )
    return {
        "heading_count": heading_count,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "figure_count": len(doc.inline_shapes),
        "page_break_count": page_break_count,
    }


def compute_docx_fidelity(generated_path: Union[str, Path], expected_path: Union[str, Path]) -> Dict[str, Any]:
    """Compare structural counts between a freshly generated DOCX and its
    expected benchmark counterpart. Returns per-category counts/diffs and
    an overall fidelity score — see module docstring for the formula."""
    generated = _counts(generated_path)
    expected = _counts(expected_path)

    diffs = {k: generated[k] - expected[k] for k in _COUNT_CATEGORIES}
    total_abs_diff = sum(abs(d) for d in diffs.values())
    total_expected = sum(expected.values())
    fidelity = 1.0 - (total_abs_diff / max(1, total_expected))
    fidelity = max(0.0, min(1.0, fidelity))

    return {
        "generated": generated,
        "expected": expected,
        "diff": diffs,
        "fidelity": round(fidelity, 4),
    }
