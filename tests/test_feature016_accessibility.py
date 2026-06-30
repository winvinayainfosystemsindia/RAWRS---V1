"""Tests for FEATURE_016: Complete Accessibility Remediation Framework.

Covers:
  - HeadingReviewStatus model (016A)
  - GET/PATCH /headings API (016A)
  - HEADING_005 validation rule — multiple H1 (016A)
  - FootnoteReviewStatus model (016D)
  - PATCH /footnotes/{id} API (016D)
  - footnote_id assignment in footnote_detector (016D)
  - GET/PATCH /metadata API (016F)
  - META_001/META_002 validation rules (016F)
  - DOCX CoreProperties written from metadata (016F)
  - Semantic list paragraph rendering in DOCX (016C)
  - ReadingOrderStatus model (016B)
  - GET /reading-order API (016B)
  - PATCH /pages/{n}/reading-order API (016B)
  - corrected_order effect on markdown block ordering (016B)
"""

from datetime import datetime, timezone
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from src.api.jobs import Job, JobStatus, _jobs
from src.api.main import app
from src.models.contracts import (
    BoundingBox,
    Document,
    Footnote,
    Heading,
    HeadingLevel,
    HeadingReviewStatus,
    FootnoteReviewStatus,
    Metadata,
    NoteType,
    Page,
    ReadingOrderStatus,
    Severity,
    Span,
    TextBlock,
    ValidationIssue,
)
from src.pipeline.phase1_pipeline import PipelineResult
from src.validation.validator import validate_document


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture(autouse=True)
def clear_jobs():
    _jobs.clear()
    yield
    _jobs.clear()


@pytest.fixture()
def client():
    return TestClient(app)


def _make_doc(
    headings: list | None = None,
    footnotes: list | None = None,
    language: str | None = None,
    title: str | None = None,
) -> Document:
    doc = Document(
        source_pdf_path="test.pdf",
        metadata=Metadata(
            filename="test.pdf",
            page_count=2,
            language=language,
            title=title,
        ),
    )
    if headings:
        doc.headings = headings
    if footnotes:
        doc.footnotes = footnotes
    return doc


def _inject_job(doc: Document, job_id: str = "test-job-a016") -> str:
    result = PipelineResult(
        source_pdf_path="test.pdf",
        success=True,
        status=doc.processing_status,
        duration_seconds=0.1,
        document=doc,
    )
    job = Job(
        job_id=job_id,
        filename="test.pdf",
        pdf_path=Path("test.pdf"),
        status=JobStatus.COMPLETE,
        created_at=datetime.now(timezone.utc),
        result=result,
    )
    _jobs[job_id] = job
    return job_id


def _h(level: int, text: str, order: int, page: int = 1, page_marker: bool = False) -> Heading:
    return Heading(level=HeadingLevel(level), text=text, page_number=page, document_order=order, is_page_marker=page_marker)


def _fn(number: int, body: str = "Note body.", anchor_page: int = 1, fn_id: str | None = None) -> Footnote:
    note = Footnote(
        note_type=NoteType.FOOTNOTE,
        number=number,
        marker=str(number),
        anchor_page_number=anchor_page,
        anchor_text=f"text with marker {number}",
        body=body,
        body_page_number=anchor_page,
        body_source_text=f"{number} {body}",
    )
    note.footnote_id = fn_id or f"fn-{number - 1}"
    return note


# ===========================================================================
# FEATURE_016A: Heading model
# ===========================================================================


class TestHeadingReviewStatusModel:
    def test_default_status_is_detected(self):
        h = _h(1, "Introduction", 0)
        assert h.review_status == HeadingReviewStatus.DETECTED

    def test_review_status_can_be_set(self):
        h = _h(1, "Introduction", 0)
        h.review_status = HeadingReviewStatus.APPROVED
        assert h.review_status == HeadingReviewStatus.APPROVED

    def test_reviewer_note_defaults_to_none(self):
        h = _h(1, "Introduction", 0)
        assert h.reviewer_note is None

    def test_all_statuses_accepted(self):
        for status in HeadingReviewStatus:
            h = Heading(
                level=HeadingLevel.H2,
                text="Test",
                page_number=1,
                document_order=0,
                review_status=status,
            )
            assert h.review_status == status

    def test_page_marker_h6_still_accepted(self):
        h = Heading(
            level=HeadingLevel.H6,
            text="5",
            page_number=5,
            document_order=10,
            is_page_marker=True,
        )
        assert h.review_status == HeadingReviewStatus.DETECTED


# ===========================================================================
# FEATURE_016A: GET /headings
# ===========================================================================


class TestGetHeadingsEndpoint:
    def test_returns_empty_when_no_headings(self, client):
        job_id = _inject_job(_make_doc())
        resp = client.get(f"/api/documents/{job_id}/headings")
        assert resp.status_code == 200
        assert resp.json()["headings"] == []

    def test_returns_content_headings_only(self, client):
        h1 = _h(1, "Title", 0)
        h2 = _h(2, "Chapter", 1)
        pm = _h(6, "1", 2, page_marker=True)
        job_id = _inject_job(_make_doc(headings=[h1, h2, pm]))
        resp = client.get(f"/api/documents/{job_id}/headings")
        headings = resp.json()["headings"]
        assert len(headings) == 2
        levels = {h["level"] for h in headings}
        assert 6 not in levels

    def test_heading_fields_present(self, client):
        job_id = _inject_job(_make_doc(headings=[_h(1, "My Title", 0)]))
        resp = client.get(f"/api/documents/{job_id}/headings")
        h = resp.json()["headings"][0]
        assert h["level"] == 1
        assert h["text"] == "My Title"
        assert h["review_status"] == "detected"
        assert h["is_page_marker"] is False

    def test_404_on_missing_job(self, client):
        resp = client.get("/api/documents/no-such-job/headings")
        assert resp.status_code == 404

    def test_409_while_processing(self, client):
        doc = _make_doc()
        job = Job(
            job_id="processing-job",
            filename="test.pdf",
            pdf_path=Path("test.pdf"),
            status=JobStatus.PROCESSING,
            created_at=datetime.now(timezone.utc),
        )
        _jobs["processing-job"] = job
        resp = client.get("/api/documents/processing-job/headings")
        assert resp.status_code == 409


# ===========================================================================
# FEATURE_016A: PATCH /headings/{document_order}
# ===========================================================================


class TestPatchHeadingEndpoint:
    def test_approve_sets_status(self, client):
        job_id = _inject_job(_make_doc(headings=[_h(2, "Chapter 1", 0)]))
        resp = client.patch(
            f"/api/documents/{job_id}/headings/0",
            json={"action": "approve"},
        )
        assert resp.status_code == 200
        assert resp.json()["review_status"] == "approved"

    def test_reject_sets_status(self, client):
        job_id = _inject_job(_make_doc(headings=[_h(2, "Not a heading", 0)]))
        resp = client.patch(
            f"/api/documents/{job_id}/headings/0",
            json={"action": "reject", "reviewer_note": "Bold body text, not a heading"},
        )
        assert resp.status_code == 200
        assert resp.json()["review_status"] == "rejected"
        assert resp.json()["reviewer_note"] == "Bold body text, not a heading"

    def test_change_level(self, client):
        job_id = _inject_job(_make_doc(headings=[_h(1, "Chapter", 0)]))
        resp = client.patch(
            f"/api/documents/{job_id}/headings/0",
            json={"level": 2},
        )
        assert resp.status_code == 200
        assert resp.json()["level"] == 2
        assert resp.json()["review_status"] == "level_changed"

    def test_change_text(self, client):
        job_id = _inject_job(_make_doc(headings=[_h(2, "Intrduction", 0)]))
        resp = client.patch(
            f"/api/documents/{job_id}/headings/0",
            json={"text": "Introduction"},
        )
        assert resp.status_code == 200
        assert resp.json()["text"] == "Introduction"

    def test_level_6_rejected(self, client):
        job_id = _inject_job(_make_doc(headings=[_h(2, "Chapter", 0)]))
        resp = client.patch(
            f"/api/documents/{job_id}/headings/0",
            json={"level": 6},
        )
        assert resp.status_code == 422

    def test_blank_text_rejected(self, client):
        job_id = _inject_job(_make_doc(headings=[_h(2, "Chapter", 0)]))
        resp = client.patch(
            f"/api/documents/{job_id}/headings/0",
            json={"text": "   "},
        )
        assert resp.status_code == 422

    def test_unknown_action_rejected(self, client):
        job_id = _inject_job(_make_doc(headings=[_h(2, "Chapter", 0)]))
        resp = client.patch(
            f"/api/documents/{job_id}/headings/0",
            json={"action": "mark_decorative"},
        )
        assert resp.status_code == 422

    def test_404_on_missing_heading(self, client):
        job_id = _inject_job(_make_doc(headings=[_h(2, "Chapter", 5)]))
        resp = client.patch(
            f"/api/documents/{job_id}/headings/99",
            json={"action": "approve"},
        )
        assert resp.status_code == 404

    def test_page_marker_not_reviewable(self, client):
        pm = _h(6, "3", 3, page_marker=True)
        job_id = _inject_job(_make_doc(headings=[pm]))
        resp = client.patch(
            f"/api/documents/{job_id}/headings/3",
            json={"action": "approve"},
        )
        assert resp.status_code == 404


# ===========================================================================
# FEATURE_016A: HEADING_005 validation rule
# ===========================================================================


class TestHeading005MultipleH1:
    def _validate(self, headings):
        doc = _make_doc(headings=headings)
        return validate_document(doc)

    def test_single_h1_no_issue(self):
        issues = self._validate([_h(1, "Title", 0), _h(2, "Chapter", 1)])
        rule_ids = {i.rule_id for i in issues}
        assert "HEADING_005" not in rule_ids

    def test_two_h1_fires_warning(self):
        issues = self._validate([_h(1, "Title A", 0), _h(1, "Title B", 1)])
        h005 = [i for i in issues if i.rule_id == "HEADING_005"]
        assert len(h005) == 1
        assert h005[0].severity == Severity.WARNING

    def test_three_h1_fires_warning_with_count(self):
        issues = self._validate([
            _h(1, "A", 0), _h(1, "B", 1), _h(1, "C", 2)
        ])
        h005 = [i for i in issues if i.rule_id == "HEADING_005"]
        assert len(h005) == 1
        assert "3" in h005[0].message

    def test_page_markers_not_counted(self):
        # H6 page markers should not affect this rule
        issues = self._validate([
            _h(1, "Title", 0),
            _h(6, "1", 1, page_marker=True),
            _h(6, "2", 2, page_marker=True),
        ])
        rule_ids = {i.rule_id for i in issues}
        assert "HEADING_005" not in rule_ids

    def test_no_headings_no_h005(self):
        issues = self._validate([])
        rule_ids = {i.rule_id for i in issues}
        assert "HEADING_005" not in rule_ids


# ===========================================================================
# FEATURE_016D: FootnoteReviewStatus model
# ===========================================================================


class TestFootnoteReviewStatusModel:
    def test_default_status_is_detected(self):
        note = _fn(1)
        assert note.review_status == FootnoteReviewStatus.DETECTED

    def test_footnote_id_assignable(self):
        note = _fn(1, fn_id="fn-0")
        assert note.footnote_id == "fn-0"

    def test_reviewer_note_defaults_to_none(self):
        note = _fn(1)
        assert note.reviewer_note is None


# ===========================================================================
# FEATURE_016D: PATCH /footnotes/{footnote_id}
# ===========================================================================


class TestPatchFootnoteEndpoint:
    def test_approve_sets_status(self, client):
        job_id = _inject_job(_make_doc(footnotes=[_fn(1, fn_id="fn-0")]))
        resp = client.patch(
            f"/api/documents/{job_id}/footnotes/fn-0",
            json={"action": "approve"},
        )
        assert resp.status_code == 200
        assert resp.json()["review_status"] == "approved"

    def test_reject_sets_status(self, client):
        job_id = _inject_job(_make_doc(footnotes=[_fn(1, fn_id="fn-0")]))
        resp = client.patch(
            f"/api/documents/{job_id}/footnotes/fn-0",
            json={"action": "reject", "reviewer_note": "False positive"},
        )
        assert resp.status_code == 200
        assert resp.json()["review_status"] == "rejected"

    def test_edit_body_sets_edited_status(self, client):
        job_id = _inject_job(_make_doc(footnotes=[_fn(1, "Orginal text.", fn_id="fn-0")]))
        resp = client.patch(
            f"/api/documents/{job_id}/footnotes/fn-0",
            json={"body": "Original text."},
        )
        assert resp.status_code == 200
        assert resp.json()["body"] == "Original text."
        assert resp.json()["review_status"] == "edited"

    def test_blank_body_rejected(self, client):
        job_id = _inject_job(_make_doc(footnotes=[_fn(1, fn_id="fn-0")]))
        resp = client.patch(
            f"/api/documents/{job_id}/footnotes/fn-0",
            json={"body": "   "},
        )
        assert resp.status_code == 422

    def test_404_on_unknown_footnote_id(self, client):
        job_id = _inject_job(_make_doc(footnotes=[_fn(1, fn_id="fn-0")]))
        resp = client.patch(
            f"/api/documents/{job_id}/footnotes/fn-999",
            json={"action": "approve"},
        )
        assert resp.status_code == 404

    def test_get_footnotes_includes_review_fields(self, client):
        job_id = _inject_job(_make_doc(footnotes=[_fn(1, fn_id="fn-0")]))
        resp = client.get(f"/api/documents/{job_id}/footnotes")
        assert resp.status_code == 200
        notes = resp.json()["footnotes"]
        assert len(notes) == 1
        assert notes[0]["footnote_id"] == "fn-0"
        assert notes[0]["review_status"] == "detected"


# ===========================================================================
# FEATURE_016F: GET/PATCH /metadata
# ===========================================================================


class TestMetadataEndpoint:
    def test_get_metadata_returns_fields(self, client):
        job_id = _inject_job(_make_doc(language="en-AU", title="Test Doc"))
        resp = client.get(f"/api/documents/{job_id}/metadata")
        assert resp.status_code == 200
        data = resp.json()
        assert data["language"] == "en-AU"
        assert data["title"] == "Test Doc"
        assert data["filename"] == "test.pdf"

    def test_get_metadata_null_when_not_set(self, client):
        job_id = _inject_job(_make_doc())
        resp = client.get(f"/api/documents/{job_id}/metadata")
        assert resp.status_code == 200
        assert resp.json()["language"] is None
        assert resp.json()["title"] is None

    def test_patch_sets_language_and_title(self, client):
        job_id = _inject_job(_make_doc())
        resp = client.patch(
            f"/api/documents/{job_id}/metadata",
            json={"language": "en-US", "title": "Annual Report"},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["language"] == "en-US"
        assert data["title"] == "Annual Report"

    def test_patch_clears_with_empty_string(self, client):
        job_id = _inject_job(_make_doc(language="en-US"))
        resp = client.patch(
            f"/api/documents/{job_id}/metadata",
            json={"language": ""},
        )
        assert resp.status_code == 200
        assert resp.json()["language"] is None

    def test_patch_author_and_subject(self, client):
        job_id = _inject_job(_make_doc())
        resp = client.patch(
            f"/api/documents/{job_id}/metadata",
            json={"author": "WinVinaya Foundation", "subject": "Accessibility"},
        )
        assert resp.status_code == 200
        assert resp.json()["author"] == "WinVinaya Foundation"
        assert resp.json()["subject"] == "Accessibility"

    def test_404_on_missing_job(self, client):
        resp = client.get("/api/documents/no-job/metadata")
        assert resp.status_code == 404


# ===========================================================================
# FEATURE_016F: META_001/META_002 validation rules
# ===========================================================================


class TestMeta001002ValidationRules:
    def test_meta001_fires_when_no_language(self):
        doc = _make_doc()
        issues = validate_document(doc)
        meta001 = [i for i in issues if i.rule_id == "META_001"]
        assert len(meta001) == 1
        assert meta001[0].severity == Severity.INFO

    def test_meta002_fires_when_no_title(self):
        doc = _make_doc()
        issues = validate_document(doc)
        meta002 = [i for i in issues if i.rule_id == "META_002"]
        assert len(meta002) == 1
        assert meta002[0].severity == Severity.INFO

    def test_no_meta001_when_language_set(self):
        doc = _make_doc(language="en-US")
        issues = validate_document(doc)
        assert not any(i.rule_id == "META_001" for i in issues)

    def test_no_meta002_when_title_set(self):
        doc = _make_doc(title="Annual Report")
        issues = validate_document(doc)
        assert not any(i.rule_id == "META_002" for i in issues)

    def test_both_clear_when_both_set(self):
        doc = _make_doc(language="fr-FR", title="Rapport Annuel")
        issues = validate_document(doc)
        meta_ids = {i.rule_id for i in issues if i.rule_id.startswith("META_")}
        assert meta_ids == set()


# ===========================================================================
# FEATURE_016F: DOCX CoreProperties written from metadata
# ===========================================================================


class TestDocxCoreProperties:
    def test_language_written_to_docx(self, tmp_path):
        from src.docx.docx_generator import generate_docx
        from docx import Document as DocxDocument

        doc = _make_doc(language="en-AU", title="Test Document")
        docx_path = tmp_path / "out.docx"
        generate_docx(doc, "# Test Document\n\nBody text.", output_path=docx_path)

        result = DocxDocument(str(docx_path))
        assert result.core_properties.language == "en-AU"
        assert result.core_properties.title == "Test Document"

    def test_no_properties_written_when_not_set(self, tmp_path):
        from src.docx.docx_generator import generate_docx
        from docx import Document as DocxDocument

        doc = _make_doc()
        docx_path = tmp_path / "out.docx"
        generate_docx(doc, "# Test\n\nBody.", output_path=docx_path)

        result = DocxDocument(str(docx_path))
        # Default python-docx CoreProperties values — should not be overwritten with None
        assert result.core_properties.title in ("", None)
        assert result.core_properties.language in ("", None)


# ===========================================================================
# FEATURE_016C: Semantic list paragraph rendering in DOCX
# ===========================================================================


def _gen_docx(tmp_path, markdown: str):
    """Generate a DOCX from the given markdown string and return the parsed document."""
    from src.docx.docx_generator import generate_docx
    from docx import Document as DocxDocument

    doc = _make_doc(language="en-US", title="Test")
    out = tmp_path / "out.docx"
    generate_docx(doc, markdown, output_path=out)
    return DocxDocument(str(out))


def _para_styles(docx_doc):
    """Return list of (style_name, text) for all non-empty paragraphs."""
    return [
        (p.style.name, p.text)
        for p in docx_doc.paragraphs
        if p.text.strip()
    ]


class TestListBulletRendering:
    def test_bullet_character_renders_as_list_bullet_style(self, tmp_path):
        result = _gen_docx(tmp_path, "• First item\n• Second item")
        styles = _para_styles(result)
        bullet_styles = [(s, t) for s, t in styles if s == "List Bullet"]
        assert len(bullet_styles) == 2
        assert bullet_styles[0][1] == "First item"
        assert bullet_styles[1][1] == "Second item"

    def test_dash_renders_as_list_bullet_style(self, tmp_path):
        result = _gen_docx(tmp_path, "- Alpha\n- Beta")
        styles = _para_styles(result)
        bullet_styles = [(s, t) for s, t in styles if s == "List Bullet"]
        assert len(bullet_styles) == 2
        assert bullet_styles[0][1] == "Alpha"

    def test_bullet_marker_stripped_from_text(self, tmp_path):
        result = _gen_docx(tmp_path, "• Item text here")
        styles = _para_styles(result)
        bullet = next((t for s, t in styles if s == "List Bullet"), None)
        assert bullet == "Item text here"
        assert not bullet.startswith("•")

    def test_arrow_bullet_recognized(self, tmp_path):
        result = _gen_docx(tmp_path, "→ Click here")
        styles = _para_styles(result)
        assert any(s == "List Bullet" for s, _ in styles)

    def test_standalone_dash_not_treated_as_list(self, tmp_path):
        # "---" separator should NOT become a list item
        result = _gen_docx(tmp_path, "---")
        styles = _para_styles(result)
        assert not any(s == "List Bullet" for s, _ in styles)

    def test_plain_paragraph_not_affected(self, tmp_path):
        result = _gen_docx(tmp_path, "This is a normal paragraph.")
        styles = _para_styles(result)
        assert not any(s in ("List Bullet", "List Number") for s, _ in styles)


class TestListNumberRendering:
    def test_numbered_item_renders_as_list_number_style(self, tmp_path):
        result = _gen_docx(tmp_path, "1. First\n2. Second\n3. Third")
        styles = _para_styles(result)
        num_styles = [(s, t) for s, t in styles if s == "List Number"]
        assert len(num_styles) == 3
        assert num_styles[0][1] == "First"
        assert num_styles[2][1] == "Third"

    def test_numbered_marker_stripped(self, tmp_path):
        result = _gen_docx(tmp_path, "1. Only item")
        styles = _para_styles(result)
        num = next((t for s, t in styles if s == "List Number"), None)
        assert num == "Only item"
        assert not num.startswith("1.")

    def test_lettered_item_renders_as_list_number(self, tmp_path):
        result = _gen_docx(tmp_path, "a. Alpha\nb. Beta")
        styles = _para_styles(result)
        num_styles = [(s, t) for s, t in styles if s == "List Number"]
        assert len(num_styles) == 2

    def test_roman_numeral_item_renders_as_list_number(self, tmp_path):
        result = _gen_docx(tmp_path, "i. First\nii. Second")
        styles = _para_styles(result)
        num_styles = [(s, t) for s, t in styles if s == "List Number"]
        assert len(num_styles) == 2

    def test_period_only_not_matched(self, tmp_path):
        # A line like ". something" has no prefix — should not match
        result = _gen_docx(tmp_path, ". Not a list")
        styles = _para_styles(result)
        assert not any(s == "List Number" for s, _ in styles)


class TestMixedListAndBodyContent:
    def test_mixed_content_preserves_styles(self, tmp_path):
        md = "Intro paragraph.\n\n• Bullet one\n• Bullet two\n\nConclusion."
        result = _gen_docx(tmp_path, md)
        styles = _para_styles(result)
        style_names = [s for s, _ in styles]
        assert "List Bullet" in style_names
        # heading from generate_docx prefix is not present here (no # heading)
        # but body paragraphs should be present
        non_list = [(s, t) for s, t in styles if s not in ("List Bullet", "List Number")]
        assert any("Intro" in t or "Conclusion" in t for _, t in non_list)

    def test_bullet_and_numbered_in_same_doc(self, tmp_path):
        md = "• Apple\n• Banana\n\n1. Step one\n2. Step two"
        result = _gen_docx(tmp_path, md)
        styles = _para_styles(result)
        assert sum(1 for s, _ in styles if s == "List Bullet") == 2
        assert sum(1 for s, _ in styles if s == "List Number") == 2


# ===========================================================================
# FEATURE_016B: Reading Order Review Workspace
# ===========================================================================


def _bbox(x0=0.0, y0=0.0, x1=100.0, y1=10.0) -> BoundingBox:
    return BoundingBox(x0=x0, y0=y0, x1=x1, y1=y1)


def _block(text: str, order: int, page: int = 1, corrected_order=None) -> TextBlock:
    return TextBlock(
        page_number=page,
        text=text,
        bbox=_bbox(y0=float(order * 12)),
        order=order,
        corrected_order=corrected_order,
    )


def _make_doc_with_blocks(*blocks: TextBlock) -> Document:
    pages_set = sorted({b.page_number for b in blocks})
    doc = Document(
        source_pdf_path="test.pdf",
        metadata=Metadata(filename="test.pdf", page_count=len(pages_set)),
    )
    for pn in pages_set:
        doc.pages.append(Page(page_number=pn))
    doc.blocks.extend(blocks)
    return doc


def _inject_doc_job(doc: Document, job_id: str = "ro-job") -> str:
    result = PipelineResult(
        source_pdf_path="test.pdf",
        success=True,
        status=doc.processing_status,
        duration_seconds=0.1,
        document=doc,
    )
    from src.api.jobs import Job, JobStatus, _jobs
    from src.api.main import app
    job = Job(
        job_id=job_id,
        filename="test.pdf",
        pdf_path=Path("test.pdf"),
        status=JobStatus.COMPLETE,
        created_at=datetime.now(timezone.utc),
        result=result,
    )
    _jobs[job_id] = job
    return job_id


class TestReadingOrderStatus:
    def test_page_defaults_to_unreviewed(self):
        page = Page(page_number=1)
        assert page.reading_order_status == ReadingOrderStatus.UNREVIEWED

    def test_text_block_defaults_corrected_order_none(self):
        block = _block("hello", order=0)
        assert block.corrected_order is None

    def test_corrected_order_set_explicitly(self):
        block = _block("hello", order=3, corrected_order=0)
        assert block.corrected_order == 0


class TestGetReadingOrderApi:
    def test_empty_when_no_page003_issues(self, client):
        doc = _make_doc_with_blocks(_block("text", 0))
        _inject_doc_job(doc, "ro-empty")
        resp = client.get("/api/documents/ro-empty/reading-order")
        assert resp.status_code == 200
        assert resp.json()["pages"] == []

    def test_includes_page003_flagged_pages(self, client):
        doc = _make_doc_with_blocks(_block("A", 0, page=2), _block("B", 1, page=2))
        doc.validation_issues.append(
            ValidationIssue(severity=Severity.WARNING, rule_id="PAGE_003", message="anomaly", page_number=2)
        )
        _inject_doc_job(doc, "ro-p3")
        resp = client.get("/api/documents/ro-p3/reading-order")
        assert resp.status_code == 200
        pages = resp.json()["pages"]
        assert len(pages) == 1
        assert pages[0]["page_number"] == 2
        assert pages[0]["reading_order_status"] == "unreviewed"
        assert len(pages[0]["blocks"]) == 2

    def test_blocks_sorted_by_effective_order(self, client):
        # block_order=1 has corrected_order=0, so it should come first
        b0 = _block("Second in PyMuPDF", order=0)
        b1 = _block("First in PyMuPDF", order=1, corrected_order=0)
        b0.corrected_order = 1
        doc = _make_doc_with_blocks(b0, b1)
        doc.validation_issues.append(
            ValidationIssue(severity=Severity.WARNING, rule_id="PAGE_003", message="anomaly", page_number=1)
        )
        _inject_doc_job(doc, "ro-sort")
        resp = client.get("/api/documents/ro-sort/reading-order")
        assert resp.status_code == 200
        blocks = resp.json()["pages"][0]["blocks"]
        assert blocks[0]["block_order"] == 1   # corrected_order=0 → first
        assert blocks[1]["block_order"] == 0   # corrected_order=1 → second

    def test_includes_already_reviewed_page(self, client):
        doc = _make_doc_with_blocks(_block("text", 0))
        doc.pages[0].reading_order_status = ReadingOrderStatus.APPROVED
        _inject_doc_job(doc, "ro-reviewed")
        resp = client.get("/api/documents/ro-reviewed/reading-order")
        assert resp.status_code == 200
        pages = resp.json()["pages"]
        assert len(pages) == 1
        assert pages[0]["reading_order_status"] == "approved"


class TestPatchReadingOrderApi:
    def test_approve_sets_status_approved(self, client):
        doc = _make_doc_with_blocks(_block("A", 0))
        doc.validation_issues.append(
            ValidationIssue(severity=Severity.WARNING, rule_id="PAGE_003", message="anomaly", page_number=1)
        )
        _inject_doc_job(doc, "ro-approve")
        resp = client.patch(
            "/api/documents/ro-approve/pages/1/reading-order",
            json={"action": "approve"},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["reading_order_status"] == "approved"

    def test_reorder_sets_corrected_order_and_status(self, client):
        b0 = _block("First in PDF", order=0)
        b1 = _block("Second in PDF", order=1)
        doc = _make_doc_with_blocks(b0, b1)
        doc.validation_issues.append(
            ValidationIssue(severity=Severity.WARNING, rule_id="PAGE_003", message="anomaly", page_number=1)
        )
        _inject_doc_job(doc, "ro-reorder")
        # Swap: put block_order=1 first, then block_order=0
        resp = client.patch(
            "/api/documents/ro-reorder/pages/1/reading-order",
            json={"action": "reorder", "block_sequence": [1, 0]},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["reading_order_status"] == "corrected"
        # Response blocks should now be in new order: block_order 1 first
        assert data["blocks"][0]["block_order"] == 1
        assert data["blocks"][1]["block_order"] == 0

    def test_reorder_without_sequence_is_422(self, client):
        doc = _make_doc_with_blocks(_block("A", 0))
        doc.validation_issues.append(
            ValidationIssue(severity=Severity.WARNING, rule_id="PAGE_003", message="anomaly", page_number=1)
        )
        _inject_doc_job(doc, "ro-bad")
        resp = client.patch(
            "/api/documents/ro-bad/pages/1/reading-order",
            json={"action": "reorder"},
        )
        assert resp.status_code == 422

    def test_unknown_action_is_422(self, client):
        doc = _make_doc_with_blocks(_block("A", 0))
        _inject_doc_job(doc, "ro-unk")
        resp = client.patch(
            "/api/documents/ro-unk/pages/1/reading-order",
            json={"action": "unknown_action"},
        )
        assert resp.status_code == 422

    def test_page_not_found_is_404(self, client):
        doc = _make_doc_with_blocks(_block("A", 0, page=1))
        _inject_doc_job(doc, "ro-404")
        resp = client.patch(
            "/api/documents/ro-404/pages/99/reading-order",
            json={"action": "approve"},
        )
        assert resp.status_code == 404


class TestCorrectedOrderAffectsMarkdown:
    def test_corrected_order_changes_block_sequence_in_output(self):
        from src.markdown.markdown_builder import build_markdown

        # Two blocks: natural order is A(0) then B(1)
        # Correct order swaps them: B should come first
        b_a = _block("Block A text", order=0, corrected_order=1)
        b_b = _block("Block B text", order=1, corrected_order=0)
        doc = _make_doc_with_blocks(b_a, b_b)
        doc.pages[0].raw_text = "Block A text\nBlock B text"
        doc.pages[0].cleaned_text = "Block A text\nBlock B text"

        markdown = build_markdown(doc)
        pos_a = markdown.find("Block A text")
        pos_b = markdown.find("Block B text")
        assert pos_b < pos_a, "Block B (corrected_order=0) should appear before Block A in output"

    def test_no_corrected_order_uses_natural_order(self):
        from src.markdown.markdown_builder import build_markdown

        b_a = _block("Alpha text", order=0)
        b_b = _block("Beta text", order=1)
        doc = _make_doc_with_blocks(b_a, b_b)
        doc.pages[0].raw_text = "Alpha text\nBeta text"
        doc.pages[0].cleaned_text = "Alpha text\nBeta text"

        markdown = build_markdown(doc)
        pos_a = markdown.find("Alpha text")
        pos_b = markdown.find("Beta text")
        assert pos_a < pos_b, "Natural order: Alpha (order=0) before Beta (order=1)"


# ===========================================================================
# FEATURE_016G: Formatting Fidelity (bold/italic preservation)
# ===========================================================================


def _span(text: str, font_flags: int, font_size: float = 12.0) -> Span:
    return Span(
        text=text,
        font_name="Times New Roman",
        font_size=font_size,
        font_flags=font_flags,
        baseline_y=10.0,
        bbox=BoundingBox(x0=0, y0=0, x1=100, y1=12),
    )


def _block_with_spans(text: str, order: int, spans: list, page: int = 1) -> TextBlock:
    b = TextBlock(
        page_number=page,
        text=text,
        bbox=_bbox(y0=float(order * 14)),
        order=order,
    )
    b.spans = spans
    return b


def _make_single_page_doc(*blocks: TextBlock, page_text: str = "") -> Document:
    doc = Document(
        source_pdf_path="test.pdf",
        metadata=Metadata(filename="test.pdf", page_count=1),
    )
    page = Page(page_number=1)
    page.raw_text = page_text
    page.cleaned_text = page_text
    doc.pages.append(page)
    doc.blocks.extend(blocks)
    return doc


class TestFormattingFidelityMarkdown:
    """016G: markdown_builder emits **...** / *...* / ***...*** based on span data."""

    def test_bold_spans_wrap_paragraph_in_double_asterisk(self):
        from src.markdown.markdown_builder import build_markdown

        blk = _block_with_spans("Bold line.", order=0, spans=[_span("Bold line.", font_flags=16)])
        doc = _make_single_page_doc(blk, page_text="Bold line.")
        md = build_markdown(doc)
        assert "**Bold line.**" in md

    def test_italic_spans_wrap_paragraph_in_single_asterisk(self):
        from src.markdown.markdown_builder import build_markdown

        blk = _block_with_spans("Italic line.", order=0, spans=[_span("Italic line.", font_flags=2)])
        doc = _make_single_page_doc(blk, page_text="Italic line.")
        md = build_markdown(doc)
        assert "*Italic line.*" in md

    def test_bold_and_italic_spans_wrap_in_triple_asterisk(self):
        from src.markdown.markdown_builder import build_markdown

        blk = _block_with_spans(
            "Bold italic.", order=0, spans=[_span("Bold italic.", font_flags=18)]  # 16+2
        )
        doc = _make_single_page_doc(blk, page_text="Bold italic.")
        md = build_markdown(doc)
        assert "***Bold italic.***" in md

    def test_mixed_bold_not_bold_produces_no_wrapping(self):
        from src.markdown.markdown_builder import build_markdown

        # One span bold, one not bold — paragraph is mixed
        blk = _block_with_spans(
            "Mixed text.", order=0,
            spans=[_span("Mixed", font_flags=16), _span(" text.", font_flags=0)],
        )
        doc = _make_single_page_doc(blk, page_text="Mixed text.")
        md = build_markdown(doc)
        assert "**Mixed text.**" not in md
        assert "Mixed text." in md

    def test_no_spans_is_bold_true_wraps_in_bold(self):
        from src.markdown.markdown_builder import build_markdown

        blk = _block("Is-bold line.", order=0)
        blk.is_bold = True
        doc = _make_single_page_doc(blk, page_text="Is-bold line.")
        md = build_markdown(doc)
        assert "**Is-bold line.**" in md

    def test_no_spans_is_bold_none_produces_no_wrapping(self):
        from src.markdown.markdown_builder import build_markdown

        blk = _block("Plain line.", order=0)  # is_bold defaults to None
        doc = _make_single_page_doc(blk, page_text="Plain line.")
        md = build_markdown(doc)
        assert "**Plain line.**" not in md
        assert "Plain line." in md

    def test_superscript_span_excluded_from_bold_check(self):
        from src.markdown.markdown_builder import build_markdown

        # Body span bold=True; superscript span (bit 1) also present — should still be bold
        blk = _block_with_spans(
            "Body text.", order=0,
            spans=[
                _span("Body text.", font_flags=16),    # bold body
                _span("1", font_flags=1),              # superscript footnote marker
            ],
        )
        doc = _make_single_page_doc(blk, page_text="Body text.")
        md = build_markdown(doc)
        assert "**Body text.**" in md

    def test_no_wrapping_when_no_blocks_for_page(self):
        from src.markdown.markdown_builder import build_markdown

        # OCR path: page has no blocks, falls back to line-by-line (no formatting)
        doc = Document(
            source_pdf_path="test.pdf",
            metadata=Metadata(filename="test.pdf", page_count=1),
        )
        page = Page(page_number=1)
        page.raw_text = "OCR line."
        page.cleaned_text = "OCR line."
        doc.pages.append(page)
        # No blocks added
        md = build_markdown(doc)
        assert "**OCR line.**" not in md
        assert "OCR line." in md


class TestFormattingFidelityDocx:
    """016G: docx_generator renders **...** / *...* / ***...*** as bold/italic runs."""

    def _runs_for_text(self, docx_doc, search_text: str):
        """Return runs from the paragraph whose full text contains search_text."""
        for para in docx_doc.paragraphs:
            if search_text in para.text:
                return para.runs
        return []

    def test_bold_markers_produce_bold_run(self, tmp_path):
        result = _gen_docx(tmp_path, "**Bold paragraph.**")
        runs = self._runs_for_text(result, "Bold paragraph.")
        assert runs, "Expected paragraph with 'Bold paragraph.'"
        assert any(r.bold for r in runs), "At least one run should be bold"

    def test_italic_markers_produce_italic_run(self, tmp_path):
        result = _gen_docx(tmp_path, "*Italic paragraph.*")
        # Note: *Italic paragraph.* is handled as body text (no image pending → no caption match)
        runs = self._runs_for_text(result, "Italic paragraph.")
        assert runs, "Expected paragraph with 'Italic paragraph.'"
        assert any(r.italic for r in runs), "At least one run should be italic"

    def test_bold_italic_markers_produce_bold_and_italic_run(self, tmp_path):
        result = _gen_docx(tmp_path, "***Bold and italic.***")
        runs = self._runs_for_text(result, "Bold and italic.")
        assert runs, "Expected paragraph with 'Bold and italic.'"
        assert any(r.bold for r in runs), "At least one run should be bold"
        assert any(r.italic for r in runs), "At least one run should be italic"

    def test_plain_text_has_no_bold(self, tmp_path):
        result = _gen_docx(tmp_path, "Plain paragraph.")
        runs = self._runs_for_text(result, "Plain paragraph.")
        assert runs, "Expected paragraph with 'Plain paragraph.'"
        assert not any(r.bold for r in runs), "Plain text should have no bold runs"

    def test_mixed_bold_and_plain_in_one_paragraph(self, tmp_path):
        result = _gen_docx(tmp_path, "Start **bold** end.")
        runs = self._runs_for_text(result, "bold")
        bold_runs = [r for r in runs if r.bold]
        assert bold_runs, "The 'bold' segment should be in a bold run"
        # Verify non-bold segments exist
        plain_runs = [r for r in runs if not r.bold and r.text.strip()]
        assert plain_runs, "Plain segments around the bold word should exist"

    def test_parse_inline_format_bold(self):
        from src.docx.docx_generator import _parse_inline_format
        result = _parse_inline_format("**hello**")
        assert result == [("hello", True, False)]

    def test_parse_inline_format_italic(self):
        from src.docx.docx_generator import _parse_inline_format
        result = _parse_inline_format("*hello*")
        assert result == [("hello", False, True)]

    def test_parse_inline_format_bold_italic(self):
        from src.docx.docx_generator import _parse_inline_format
        result = _parse_inline_format("***hello***")
        assert result == [("hello", True, True)]

    def test_parse_inline_format_mixed_segments(self):
        from src.docx.docx_generator import _parse_inline_format
        result = _parse_inline_format("Start **bold** and *italic* end.")
        assert result == [
            ("Start ", False, False),
            ("bold", True, False),
            (" and ", False, False),
            ("italic", False, True),
            (" end.", False, False),
        ]

    def test_parse_inline_format_plain_text(self):
        from src.docx.docx_generator import _parse_inline_format
        result = _parse_inline_format("no formatting here")
        assert result == [("no formatting here", False, False)]
