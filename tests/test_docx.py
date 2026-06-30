"""Tests for src/docx/docx_generator.py."""

from pathlib import Path
from typing import List

import zipfile

import fitz
import pytest
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph
from lxml import etree

from src.docx.docx_generator import _MAX_IMAGE_WIDTH, _add_heading, _safe_run_text, generate_docx
from src.headings.heading_detector import detect_headings
from src.markdown.markdown_builder import build_markdown
from src.models.contracts import Document, Metadata, Page
from src.ocr.extractor import extract_text
from src.parser.pdf_parser import parse_pdf

SAMPLE_PDF_DIR = Path(__file__).resolve().parents[1] / "samples" / "benchmark" / "pdfs"
SAMPLE_PDFS = sorted(SAMPLE_PDF_DIR.glob("*.pdf"))

_BLACK = RGBColor(0, 0, 0)


def _dummy_document() -> Document:
    return Document(source_pdf_path="dummy.pdf", metadata=Metadata(filename="dummy.pdf"))


def _make_png(path: Path, width: int = 50, height: int = 40, color=(255, 0, 0)) -> Path:
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, width, height))
    pix.set_rect(pix.irect, color)
    pix.save(str(path))
    return path


def _heading_paragraphs(docx_doc: DocxDocument) -> List[Paragraph]:
    return [p for p in docx_doc.paragraphs if p.style.name.startswith("Heading")]


def _paragraph_with_text(docx_doc: DocxDocument, text: str) -> Paragraph:
    for paragraph in docx_doc.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise AssertionError(f"No paragraph found with text: {text!r}")


def _count_page_breaks(docx_doc: DocxDocument) -> int:
    count = 0
    for paragraph in docx_doc.paragraphs:
        for run in paragraph.runs:
            count += len(run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type='page']"))
    return count


class TestDocxGeneration:
    def test_generates_a_readable_docx_file(self, tmp_path: Path) -> None:
        output_path = tmp_path / "out.docx"
        result_path = generate_docx(_dummy_document(), "# Title", output_path=output_path)

        assert result_path == output_path
        assert output_path.is_file()
        DocxDocument(str(output_path))  # raises if not a valid docx

    def test_default_output_path_uses_source_pdf_stem(self) -> None:
        from src.docx.docx_generator import _resolve_output_path

        document = Document(
            source_pdf_path="samples/benchmark/pdfs/Delpit (1988).pdf",
            metadata=Metadata(filename="Delpit (1988).pdf"),
        )
        resolved = _resolve_output_path(document, None)
        assert resolved == Path("outputs/docx") / "Delpit (1988).docx"

    def test_empty_markdown_still_produces_valid_docx(self, tmp_path: Path) -> None:
        output_path = tmp_path / "empty.docx"
        generate_docx(_dummy_document(), "", output_path=output_path)
        assert output_path.is_file()
        DocxDocument(str(output_path))


class TestHeadingHierarchy:
    def test_all_six_levels_map_to_correct_word_styles(self, tmp_path: Path) -> None:
        markdown = "# T1\n\n## T2\n\n### T3\n\n#### T4\n\n##### T5\n\n###### 1"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        headings = _heading_paragraphs(docx_doc)
        assert [h.style.name for h in headings] == [f"Heading {n}" for n in range(1, 7)]
        assert [h.text for h in headings] == ["T1", "T2", "T3", "T4", "T5", "1"]

    @pytest.mark.parametrize(
        "level,expected_pt", [(1, 16), (2, 14), (3, 12), (4, 12), (5, 12)]
    )
    def test_heading_font_matches_heading_rules(
        self, tmp_path: Path, level: int, expected_pt: int
    ) -> None:
        markdown = f"{'#' * level} Heading Text"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        heading = _heading_paragraphs(docx_doc)[0]
        run = heading.runs[0]

        assert run.font.name == "Times New Roman"
        assert run.font.size.pt == expected_pt
        assert run.font.bold is True
        assert run.font.color.rgb == _BLACK

    def test_h6_page_marker_inherits_style_no_explicit_run_formatting(
        self, tmp_path: Path
    ) -> None:
        # H6 page markers must NOT carry explicit run-property overrides -
        # they inherit the Heading 6 style defaults to match the benchmark
        # human-remediated DOCX convention (bare numeric text, no rpr element).
        markdown = "###### 1"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        heading = _heading_paragraphs(docx_doc)[0]
        assert heading.style.name == "Heading 6"
        for run in heading.runs:
            assert run.font.bold is None  # not explicitly overridden
            assert run.font.size is None  # not explicitly overridden
            assert run.font.name is None  # not explicitly overridden

    @pytest.mark.parametrize("level", [1, 2, 3, 4, 5])
    def test_empty_heading_text_still_gets_compliant_formatting(self, level: int) -> None:
        # add_heading() creates zero runs for empty text, which would
        # otherwise skip the font-override loop entirely and leave the
        # paragraph rendered with Word's built-in (non-compliant)
        # Heading-N theme defaults. _add_heading must guarantee a run
        # exists before applying formatting, regardless of text content.
        docx_doc = DocxDocument()
        _add_heading(docx_doc, level, "")

        heading_sizes = {1: 16, 2: 14, 3: 12, 4: 12, 5: 12}
        paragraph = docx_doc.paragraphs[0]
        assert len(paragraph.runs) >= 1
        for run in paragraph.runs:
            assert run.font.name == "Times New Roman"
            assert run.font.size.pt == heading_sizes[level]
            assert run.font.bold is True
            assert run.font.color.rgb == _BLACK


class TestNavigationPaneCompatibility:
    def test_headings_use_builtin_heading_styles(self, tmp_path: Path) -> None:
        # Word's Navigation Pane reads paragraphs styled with the
        # built-in Heading 1-9 styles directly; using add_heading()
        # with these style names is sufficient for nav pane inclusion.
        markdown = "# Document Title\n\nbody text\n\n## Introduction\n\nmore text"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        nav_eligible_styles = {f"Heading {n}" for n in range(1, 7)}
        headings = _heading_paragraphs(docx_doc)
        assert len(headings) == 2
        assert all(h.style.name in nav_eligible_styles for h in headings)

    def test_page_marker_h6_is_also_nav_pane_eligible(self, tmp_path: Path) -> None:
        markdown = "###### 1"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        heading = _heading_paragraphs(docx_doc)[0]
        assert heading.style.name == "Heading 6"


class TestPageBreakInsertion:
    def test_break_inserted_between_two_pages(self, tmp_path: Path) -> None:
        markdown = "page one text\n\n<!-- pagebreak -->\n\npage two text"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        assert _count_page_breaks(docx_doc) == 1

    def test_trailing_page_break_is_omitted(self, tmp_path: Path) -> None:
        markdown = "page one text\n\n<!-- pagebreak -->"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        assert _count_page_breaks(docx_doc) == 0

    def test_three_pages_produce_two_breaks(self, tmp_path: Path) -> None:
        markdown = "one\n\n<!-- pagebreak -->\n\ntwo\n\n<!-- pagebreak -->\n\nthree\n\n<!-- pagebreak -->"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        assert _count_page_breaks(docx_doc) == 2


class TestImageInsertion:
    def test_image_reference_is_inserted_as_inline_picture(self, tmp_path: Path) -> None:
        image_path = _make_png(tmp_path / "fig1.png")
        markdown = f"![Figure 1]({image_path.as_posix()})"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        assert len(docx_doc.inline_shapes) == 1

    def test_image_paragraph_is_center_aligned(self, tmp_path: Path) -> None:
        image_path = _make_png(tmp_path / "fig1.png")
        markdown = f"![Figure 1]({image_path.as_posix()})"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        image_paragraphs = [p for p in docx_doc.paragraphs if p.runs and p.runs[0]._element.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
        )]
        assert len(image_paragraphs) == 1
        assert image_paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_missing_image_file_is_skipped_without_raising(self, tmp_path: Path) -> None:
        markdown = "![Figure 1](does/not/exist.png)"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)  # must not raise

        docx_doc = DocxDocument(str(output_path))
        assert len(docx_doc.inline_shapes) == 0

    def test_oversized_image_is_scaled_to_max_width(self, tmp_path: Path) -> None:
        image_path = _make_png(tmp_path / "big.png", width=1000, height=500)
        markdown = f"![Figure 1]({image_path.as_posix()})"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        picture = docx_doc.inline_shapes[0]
        assert picture.width <= _MAX_IMAGE_WIDTH
        # aspect ratio (2:1) should be preserved after scaling
        assert abs((picture.width / picture.height) - 2.0) < 0.01


class TestImageAltTextMetadata:
    """Phase F.4: the inserted picture's docPr descr/title attributes -
    the OOXML accessibility metadata a screen reader reads - must come
    from whatever alt text is already in the markdown's ![alt](path)
    syntax."""

    def _docpr_for_first_image(self, docx_doc: DocxDocument):
        for paragraph in docx_doc.paragraphs:
            for run in paragraph.runs:
                inlines = run._element.findall(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
                )
                for inline in inlines:
                    return inline.find(
                        "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr"
                    )
        return None

    def test_alt_text_sets_descr_and_title(self, tmp_path: Path) -> None:
        image_path = _make_png(tmp_path / "fig1.png")
        alt_text = "Figure 1: A diagram of the process.: description pending human review"
        markdown = f"![{alt_text}]({image_path.as_posix()})"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        doc_pr = self._docpr_for_first_image(docx_doc)
        assert doc_pr is not None
        assert doc_pr.get("descr") == alt_text
        assert doc_pr.get("title") == alt_text

    def test_empty_alt_text_leaves_descr_unset(self, tmp_path: Path) -> None:
        image_path = _make_png(tmp_path / "fig1.png")
        markdown = f"![]({image_path.as_posix()})"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        doc_pr = self._docpr_for_first_image(docx_doc)
        assert doc_pr is not None
        assert doc_pr.get("descr") is None
        assert doc_pr.get("title") is None


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_W_URI = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _read_footnotes_xml(docx_path: Path):
    """Return the parsed footnotes.xml element tree from a saved DOCX,
    or None if the part is absent."""
    with zipfile.ZipFile(str(docx_path)) as zf:
        if "word/footnotes.xml" not in zf.namelist():
            return None
        return etree.fromstring(zf.read("word/footnotes.xml"))


def _footnote_elements(footnotes_root) -> list:
    """User footnote elements (id >= 1) from footnotes.xml."""
    return [
        el for el in footnotes_root.findall(f"{_W_NS}footnote")
        if int(el.get(f"{_W_NS}id", "-99")) >= 1
    ]


class TestFootnoteEndnoteRendering:
    """Phase K: a markdown ``[^label]`` inline reference plus its
    matching ``[^label]: body`` definition must produce native OOXML
    footnotes — a ``w:footnoteReference`` run in the document body
    linked by id to a ``w:footnote`` element in ``word/footnotes.xml``.
    """

    def _footnote_references(self, paragraph):
        return paragraph._p.findall(f".//{_W_NS}footnoteReference")

    def _hyperlinks(self, paragraph):
        return paragraph._p.findall(f".//{_W_NS}hyperlink")

    def test_inline_reference_renders_as_native_footnote_reference(
        self, tmp_path: Path
    ) -> None:
        markdown = "A claim with a marker[^p1-1]."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        # paragraph.text does not include w:footnoteReference content
        paragraph = _paragraph_with_text(docx_doc, "A claim with a marker.")
        refs = self._footnote_references(paragraph)
        assert len(refs) == 1

    def test_inline_reference_is_superscript(self, tmp_path: Path) -> None:
        markdown = "A claim with a marker[^p1-1]."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        paragraph = _paragraph_with_text(docx_doc, "A claim with a marker.")
        ref = self._footnote_references(paragraph)[0]
        run = ref.getparent()
        vert_align = run.find(f"{_W_NS}rPr/{_W_NS}vertAlign")
        assert vert_align is not None
        assert vert_align.get(f"{_W_NS}val") == "superscript"

    def test_no_hyperlinks_in_footnote_references(self, tmp_path: Path) -> None:
        # Native footnotes must not use the old bookmark/hyperlink mechanism.
        markdown = "A claim with a marker[^p1-1]."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        paragraph = _paragraph_with_text(docx_doc, "A claim with a marker.")
        assert self._hyperlinks(paragraph) == []

    def test_footnotes_xml_part_is_present(self, tmp_path: Path) -> None:
        markdown = "A claim[^p1-1].\n\n[^p1-1]: The note body."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        assert _read_footnotes_xml(output_path) is not None

    def test_footnote_body_text_in_footnotes_xml(self, tmp_path: Path) -> None:
        markdown = "A claim[^p1-1].\n\n[^p1-1]: The note body."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        root = _read_footnotes_xml(output_path)
        user_notes = _footnote_elements(root)
        assert len(user_notes) == 1
        # Body text is in a w:t element inside the user footnote
        all_text = "".join(
            t.text or ""
            for t in user_notes[0].iter(f"{_W_NS}t")
        ).strip()
        assert "The note body." in all_text

    def test_footnote_id_matches_between_reference_and_xml(
        self, tmp_path: Path
    ) -> None:
        markdown = "A claim[^p1-1].\n\n[^p1-1]: The note body."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        paragraph = _paragraph_with_text(docx_doc, "A claim.")
        ref_id = self._footnote_references(paragraph)[0].get(f"{_W_NS}id")

        root = _read_footnotes_xml(output_path)
        xml_ids = {el.get(f"{_W_NS}id") for el in _footnote_elements(root)}
        assert ref_id in xml_ids

    def test_definition_body_not_in_main_document_paragraphs(
        self, tmp_path: Path
    ) -> None:
        # With native footnotes the body lives in footnotes.xml, not in
        # the document body — no paragraph with the note body text should
        # appear in docx_doc.paragraphs.
        markdown = "A claim[^p1-1].\n\n[^p1-1]: The note body."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        body_texts = [p.text for p in docx_doc.paragraphs]
        assert not any("The note body." in t for t in body_texts)

    def test_footnote_body_uses_smaller_font(self, tmp_path: Path) -> None:
        # 10pt body font in footnotes.xml (sz val="20" = 20 half-points).
        markdown = "A claim[^p1-1].\n\n[^p1-1]: The note body."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        root = _read_footnotes_xml(output_path)
        user_notes = _footnote_elements(root)
        sz_vals = [
            el.get(f"{_W_NS}val")
            for el in user_notes[0].iter(f"{_W_NS}sz")
        ]
        assert "20" in sz_vals  # 20 half-points == 10 pt

    def test_text_around_reference_is_preserved(self, tmp_path: Path) -> None:
        # paragraph.text concatenates w:t content only; w:footnoteReference
        # has no w:t, so the number is absent from paragraph.text.
        markdown = "Before the marker[^p1-1] and after it."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        paragraph = _paragraph_with_text(docx_doc, "Before the marker and after it.")
        assert paragraph is not None

    def test_text_with_no_reference_is_unaffected(self, tmp_path: Path) -> None:
        markdown = "Ordinary text with no footnote markers at all."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        paragraph = _paragraph_with_text(docx_doc, "Ordinary text with no footnote markers at all.")
        assert len(paragraph.runs) == 1
        assert self._hyperlinks(paragraph) == []
        assert self._footnote_references(paragraph) == []

    def test_multiple_references_have_distinct_ids(self, tmp_path: Path) -> None:
        markdown = (
            "First claim[^p1-1].\n\n"
            "Second claim[^p1-2].\n\n"
            "[^p1-1]: First note.\n\n"
            "[^p1-2]: Second note."
        )
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        root = _read_footnotes_xml(output_path)
        user_notes = _footnote_elements(root)
        ids = [el.get(f"{_W_NS}id") for el in user_notes]
        assert len(ids) == 2
        assert ids[0] != ids[1]

    def test_document_with_footnotes_still_opens_cleanly(self, tmp_path: Path) -> None:
        markdown = "A claim with a marker[^p1-1].\n\n[^p1-1]: The note body."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        DocxDocument(str(output_path))  # must not raise


class TestCaptionInsertion:
    def test_caption_after_image_is_centered_and_italic(self, tmp_path: Path) -> None:
        image_path = _make_png(tmp_path / "fig1.png")
        markdown = f"![Figure 1]({image_path.as_posix()})\n\n*A diagram of the process.*"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        caption = _paragraph_with_text(docx_doc, "A diagram of the process.")
        assert caption.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert caption.runs[0].font.italic is True

    def test_asterisk_line_without_preceding_image_is_italic_body_text(self, tmp_path: Path) -> None:
        # 016G: *text* in body position (no preceding image) is rendered as
        # italic body text — not a caption (not centered), but italic formatting
        # IS applied because _parse_inline_format interprets *...* as italic.
        markdown = "*not a caption*"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        paragraph = _paragraph_with_text(docx_doc, "not a caption")
        assert paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER  # not a caption style
        assert paragraph.runs[0].font.italic is True  # but italic formatting IS applied


class TestFontSettings:
    def test_body_paragraph_uses_times_new_roman_12pt_black_not_bold(self, tmp_path: Path) -> None:
        markdown = "Just a plain sentence of body text."
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        paragraph = _paragraph_with_text(docx_doc, "Just a plain sentence of body text.")
        run = paragraph.runs[0]
        assert run.font.name == "Times New Roman"
        assert run.font.size.pt == 12
        assert run.font.bold is False
        assert run.font.color.rgb == _BLACK

    def test_normal_style_default_font(self, tmp_path: Path) -> None:
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), "irrelevant", output_path=output_path)

        docx_doc = DocxDocument(str(output_path))
        normal = docx_doc.styles["Normal"]
        assert normal.font.name == "Times New Roman"
        assert normal.font.size.pt == 12
        assert normal.font.color.rgb == _BLACK


class TestXmlSanitizationSafetyGuard:
    """XML Sanitization Architecture, Layer 3: a last-resort guard,
    independent of Layer 1 - markdown_content is handed to generate_docx()
    directly here (bypassing src/markdown/markdown_builder.py entirely),
    which is exactly what proves this layer prevents the crash on its
    own even if every upstream sanitization step were somehow skipped.
    Reproduces the original production failure: "All strings must be
    XML compatible..." (a ValueError from lxml) at generate_docx."""

    def test_safe_run_text_strips_illegal_characters_directly(self) -> None:
        cleaned = _safe_run_text("before\x01after")
        assert cleaned == "beforeafter"

    def test_safe_run_text_is_a_no_op_on_clean_text(self) -> None:
        assert _safe_run_text("Perfectly ordinary text.") == "Perfectly ordinary text."

    def test_body_paragraph_with_control_character_does_not_crash(self, tmp_path: Path) -> None:
        markdown = "before\x01after"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)  # must not raise

        docx_doc = DocxDocument(str(output_path))
        assert _paragraph_with_text(docx_doc, "beforeafter") is not None

    def test_heading_with_control_character_does_not_crash(self, tmp_path: Path) -> None:
        markdown = "# before\x01after"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)  # must not raise

        docx_doc = DocxDocument(str(output_path))
        headings = _heading_paragraphs(docx_doc)
        assert any(h.text == "beforeafter" for h in headings)

    def test_caption_with_control_character_does_not_crash(self, tmp_path: Path) -> None:
        image_path = _make_png(tmp_path / "fig1.png")
        markdown = f"![Figure 1]({image_path.as_posix()})\n\n*before\x01after*"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)  # must not raise

        docx_doc = DocxDocument(str(output_path))
        assert _paragraph_with_text(docx_doc, "beforeafter") is not None

    def test_footnote_definition_with_control_character_does_not_crash(
        self, tmp_path: Path
    ) -> None:
        # The sanitized body text must reach footnotes.xml without crashing.
        markdown = "A claim with a marker[^p1-1].\n\n[^p1-1]: before\x01after"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)  # must not raise

        root = _read_footnotes_xml(output_path)
        assert root is not None
        all_text = "".join(t.text or "" for t in root.iter(f"{_W_NS}t"))
        assert "beforeafter" in all_text

    def test_image_alt_text_with_control_character_does_not_crash(self, tmp_path: Path) -> None:
        image_path = _make_png(tmp_path / "fig1.png")
        markdown = f"![before\x01after]({image_path.as_posix()})"
        output_path = tmp_path / "out.docx"
        generate_docx(_dummy_document(), markdown, output_path=output_path)  # must not raise

        docx_doc = DocxDocument(str(output_path))
        picture = docx_doc.inline_shapes[0]
        doc_properties = picture._inline.docPr
        assert doc_properties.get("descr") == "beforeafter"
        assert doc_properties.get("title") == "beforeafter"


_HEADING_SIZES_PT = {1: 16, 2: 14, 3: 12, 4: 12, 5: 12, 6: 12}


@pytest.mark.parametrize("sample_pdf_path", SAMPLE_PDFS, ids=[p.name for p in SAMPLE_PDFS])
class TestDocxGenerationEndToEndWithRealPdfs:
    def test_full_pipeline_produces_valid_docx(self, sample_pdf_path: Path, tmp_path: Path) -> None:
        document = parse_pdf(sample_pdf_path)
        extract_text(document)  # Phase A: populates real text for born-digital PDFs
        detect_headings(document)
        markdown_content = build_markdown(document)

        output_path = tmp_path / "out.docx"
        result_path = generate_docx(document, markdown_content, output_path=output_path)

        assert result_path.is_file()
        docx_doc = DocxDocument(str(result_path))

        # Every PDF page must have its H6 marker, regardless of how many
        # real content headings were also detected (varies per PDF: the
        # scanned PDF still gets none, since OCR isn't implemented yet).
        headings = _heading_paragraphs(docx_doc)
        page_markers = [h for h in headings if h.style.name == "Heading 6"]
        assert len(page_markers) == len(document.pages)

    def test_every_heading_complies_with_heading_rules(
        self, sample_pdf_path: Path, tmp_path: Path
    ) -> None:
        document = parse_pdf(sample_pdf_path)
        extract_text(document)
        detect_headings(document)
        markdown_content = build_markdown(document)

        output_path = tmp_path / "out.docx"
        generate_docx(document, markdown_content, output_path=output_path)
        docx_doc = DocxDocument(str(output_path))

        for heading in _heading_paragraphs(docx_doc):
            level = int(heading.style.name.split()[-1])
            assert heading.runs, f"heading paragraph has no runs: {heading.text!r}"
            for run in heading.runs:
                if level == 6:
                    # H6 page markers inherit Heading 6 style — no explicit run overrides.
                    assert run.font.bold is None
                    assert run.font.size is None
                    assert run.font.name is None
                else:
                    assert run.font.name == "Times New Roman"
                    assert run.font.size.pt == _HEADING_SIZES_PT[level]
                    assert run.font.bold is True
                    assert run.font.color.rgb == _BLACK

    def test_every_body_paragraph_complies_with_body_text_rules(
        self, sample_pdf_path: Path, tmp_path: Path
    ) -> None:
        document = parse_pdf(sample_pdf_path)
        extract_text(document)
        detect_headings(document)
        markdown_content = build_markdown(document)

        output_path = tmp_path / "out.docx"
        generate_docx(document, markdown_content, output_path=output_path)
        docx_doc = DocxDocument(str(output_path))

        body_paragraphs = [
            p
            for p in docx_doc.paragraphs
            if not p.style.name.startswith("Heading") and p.text.strip()
        ]
        for paragraph in body_paragraphs:
            for run in paragraph.runs:
                if not run.text.strip():
                    continue
                assert run.font.name == "Times New Roman"
                assert run.font.size.pt == 12
                # 016G: bold body runs are valid when source spans are uniformly bold —
                # assert bold is not None (explicitly set), not that it is False.
                assert run.font.bold in (True, False)
                assert run.font.color.rgb == _BLACK
