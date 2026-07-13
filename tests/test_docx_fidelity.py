"""Tests for src/verification/docx_fidelity.py (M-3.3)."""

from pathlib import Path

from docx import Document as DocxDocument

from src.verification.docx_fidelity import compute_docx_fidelity

_BENCHMARK_DOCX = Path("samples/benchmark/remediated_docx/1. Nature of Enquiry.docx")


def _make_docx(tmp_path: Path, name: str, headings=1, paragraphs=1, tables=0, page_breaks=0) -> Path:
    doc = DocxDocument()
    for i in range(headings):
        doc.add_heading(f"H{i}", level=1)
    for i in range(paragraphs):
        doc.add_paragraph(f"paragraph {i}")
    for _ in range(tables):
        doc.add_table(rows=1, cols=1)
    for _ in range(page_breaks):
        doc.add_page_break()
    path = tmp_path / name
    doc.save(str(path))
    return path


class TestComputeDocxFidelity:
    def test_identical_documents_score_perfect_fidelity(self, tmp_path):
        a = _make_docx(tmp_path, "a.docx", headings=2, paragraphs=3, tables=1, page_breaks=1)
        b = _make_docx(tmp_path, "b.docx", headings=2, paragraphs=3, tables=1, page_breaks=1)
        result = compute_docx_fidelity(a, b)
        assert result["fidelity"] == 1.0
        assert all(d == 0 for d in result["diff"].values())

    def test_missing_table_reduces_fidelity(self, tmp_path):
        generated = _make_docx(tmp_path, "gen.docx", headings=1, paragraphs=1, tables=0)
        expected = _make_docx(tmp_path, "exp.docx", headings=1, paragraphs=1, tables=1)
        result = compute_docx_fidelity(generated, expected)
        assert result["fidelity"] < 1.0
        assert result["diff"]["table_count"] == -1

    def test_fidelity_never_negative(self, tmp_path):
        generated = _make_docx(tmp_path, "gen.docx", headings=0, paragraphs=0, tables=5, page_breaks=5)
        expected = _make_docx(tmp_path, "exp.docx", headings=1, paragraphs=1)
        result = compute_docx_fidelity(generated, expected)
        assert result["fidelity"] >= 0.0

    def test_real_benchmark_docx_self_comparison_is_perfect(self):
        """Sanity check against a real corpus file, not just synthetic
        python-docx output — guards against a python-docx version quirk
        that a from-scratch document wouldn't expose."""
        if not _BENCHMARK_DOCX.exists():
            return
        result = compute_docx_fidelity(_BENCHMARK_DOCX, _BENCHMARK_DOCX)
        assert result["fidelity"] == 1.0


if __name__ == "__main__":
    # ponytail: minimal runnable self-check, no pytest required
    import tempfile

    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        a = _make_docx(tmp, "a.docx", headings=2, paragraphs=3, tables=1)
        b = _make_docx(tmp, "b.docx", headings=2, paragraphs=3, tables=1)
        r = compute_docx_fidelity(a, b)
        assert r["fidelity"] == 1.0, r
        c = _make_docx(tmp, "c.docx", headings=2, paragraphs=3, tables=0)
        r2 = compute_docx_fidelity(c, b)
        assert r2["fidelity"] < 1.0, r2
    print("docx_fidelity self-check OK")
