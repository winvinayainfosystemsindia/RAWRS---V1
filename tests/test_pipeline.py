"""Tests for src/pipeline/phase1_pipeline.py.

All existing calls below pass enable_ocr=False: none of these tests are
about OCR itself, and Docling's full-page OCR pipeline (Phase D.1) is
slow (roughly 1-3 minutes per OCR_REQUIRED page) - leaving real OCR
enabled here would make every one of these run for many minutes against
O'Leary's 13 OCR_REQUIRED pages and could change its long-assumed
"empty text" fixture behavior some assertions below depend on (e.g.
"DOC_001 does not fire", "HEADING_002 fires"). The dedicated Docling
integration tests live in tests/test_docling_engine.py and the
TestDoclingIntegration class below; the dedicated Surya fallback
integration tests live in tests/test_surya_engine.py and the
TestSuryaFallbackWiring / TestSuryaIntegration classes below; the
dedicated Structure Detection (Phase H) tests live in
tests/test_structure_detector.py and the TestStructureDetectionWiring /
TestStructureDetectionDoesNotChangeExistingOutputs classes below; the
dedicated Footnote/Endnote Detection (Phase K) tests live in
tests/test_footnote_detector.py and the TestFootnoteDetectionWiring /
TestFootnoteDetectionBenchmarkRegression classes below.
"""

import json
from pathlib import Path

import fitz
import pytest

from conftest import a_scanned_pdf
from src.models.contracts import ExtractionMethod, OCRConfidence, PageType, ProcessingStatus
from src.pipeline.phase1_pipeline import run_pipeline

SAMPLE_PDF_DIR = Path(__file__).resolve().parents[1] / "samples" / "benchmark" / "pdfs"
SAMPLE_PDFS = sorted(SAMPLE_PDF_DIR.glob("*.pdf"))
A_SAMPLE_PDF = a_scanned_pdf()  # manifest-declared scanned/OCR-required PDF
# Pinned by name, not by capability lookup: several tests below assert
# Calderhead-specific content (image counts, dataset paths) - per the
# Benchmark Infrastructure Audit, fixture-document tests with
# content-specific assertions should name their file explicitly rather
# than resolve to "whichever born-digital PDF sorts first".
A_DIGITAL_SAMPLE_PDF = SAMPLE_PDF_DIR / "5.Teachingas a profession_Calderhead.pdf"


@pytest.mark.parametrize("sample_pdf_path", SAMPLE_PDFS, ids=[p.name for p in SAMPLE_PDFS])
class TestSuccessfulProcessing:
    def test_pipeline_succeeds_for_every_sample_pdf(self, sample_pdf_path: Path, tmp_path: Path) -> None:
        result = run_pipeline(sample_pdf_path, output_root=tmp_path, enable_ocr=False)

        assert result.success is True
        assert result.status == ProcessingStatus.VALIDATED
        assert result.failed_stage is None
        assert result.error_message is None
        assert result.document is not None
        assert result.document.processing_status == ProcessingStatus.VALIDATED

    def test_duration_is_recorded(self, sample_pdf_path: Path, tmp_path: Path) -> None:
        result = run_pipeline(sample_pdf_path, output_root=tmp_path, enable_ocr=False)
        assert result.duration_seconds > 0


def _build_dirty_pdf(path: Path) -> None:
    """A real PDF whose text layer contains a control character - the
    exact production failure mode (a broken font ToUnicode mapping can
    legitimately decode a glyph to one; see the XML Sanitization
    Architecture Review, docs/DECISIONS_LOG.md, for the empirical
    verification that this survives PyMuPDF's own
    insert_text -> save -> get_text round-trip)."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "A heading\x01here", fontname="helv", fontsize=16)
    page.insert_text((72, 100), "Ordinary body text on the same dirty\x02 page.")
    doc.save(str(path))
    doc.close()


class TestXmlSanitizationEndToEnd:
    """Reproduces the original production failure end-to-end: before
    the XML Sanitization Architecture existed, this exact fixture
    crashed run_pipeline at the generate_docx stage with "All strings
    must be XML compatible..." (a ValueError from lxml)."""

    def test_pipeline_succeeds_on_a_pdf_with_xml_illegal_characters(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "dirty.pdf"
        _build_dirty_pdf(pdf_path)

        result = run_pipeline(pdf_path, output_root=tmp_path / "out", enable_ocr=False)

        assert result.success is True
        assert result.failed_stage is None
        assert result.error_message is None
        assert result.docx_path is not None
        assert result.docx_path.is_file()

    def test_sanitization_events_are_recorded_and_surfaced_as_doc_004(
        self, tmp_path: Path
    ) -> None:
        pdf_path = tmp_path / "dirty.pdf"
        _build_dirty_pdf(pdf_path)

        result = run_pipeline(pdf_path, output_root=tmp_path / "out", enable_ocr=False)

        assert len(result.document.sanitization_events) > 0
        doc_004_issues = [i for i in result.validation_issues if i.rule_id == "DOC_004"]
        assert len(doc_004_issues) > 0
        assert all(i.severity.value == "warning" for i in doc_004_issues)

    def test_generated_docx_contains_no_control_characters(self, tmp_path: Path) -> None:
        from docx import Document as DocxDocument

        pdf_path = tmp_path / "dirty.pdf"
        _build_dirty_pdf(pdf_path)

        result = run_pipeline(pdf_path, output_root=tmp_path / "out", enable_ocr=False)

        docx_doc = DocxDocument(str(result.docx_path))
        full_text = "\n".join(p.text for p in docx_doc.paragraphs)
        assert "\x01" not in full_text
        assert "\x02" not in full_text


class TestMissingPdf:
    def test_missing_pdf_fails_gracefully_at_parse_stage(self, tmp_path: Path) -> None:
        missing_path = tmp_path / "does_not_exist.pdf"
        result = run_pipeline(missing_path, output_root=tmp_path / "out", enable_ocr=False)

        assert result.success is False
        assert result.status == ProcessingStatus.FAILED
        assert result.failed_stage == "parse_pdf"
        assert result.error_message is not None
        assert result.document is None
        assert result.markdown_path is None
        assert result.docx_path is None
        assert result.report_path is None

    def test_missing_pdf_does_not_raise(self, tmp_path: Path) -> None:
        missing_path = tmp_path / "nope.pdf"
        # must not raise - "fail gracefully" means a structured result,
        # never an uncaught exception
        run_pipeline(missing_path, output_root=tmp_path / "out", enable_ocr=False)


class TestStageFailureHandling:
    def test_image_extraction_failure_is_caught_and_reported(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        def _boom(document, output_dir):
            raise RuntimeError("simulated image extraction failure")

        monkeypatch.setattr("src.pipeline.phase1_pipeline.extract_images", _boom)

        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        assert result.success is False
        assert result.failed_stage == "extract_images"
        assert "simulated image extraction failure" in result.error_message
        assert result.document is not None
        assert result.document.processing_status == ProcessingStatus.FAILED
        # earlier stage's output (none here, parsing has no file output) stays None
        assert result.markdown_path is None
        assert result.docx_path is None

    def test_docx_generation_failure_preserves_markdown_output(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        def _boom(document, markdown_content, output_path=None):
            raise RuntimeError("simulated docx generation failure")

        monkeypatch.setattr("src.pipeline.phase1_pipeline.generate_docx", _boom)

        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        assert result.success is False
        assert result.failed_stage == "generate_docx"
        assert result.document.processing_status == ProcessingStatus.FAILED
        # markdown stage ran and wrote its file successfully before the failure
        assert result.markdown_path is not None
        assert result.markdown_path.is_file()
        assert result.docx_path is None
        assert result.report_path is None

    def test_validation_failure_preserves_markdown_and_docx_output(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        def _boom(document):
            raise RuntimeError("simulated validation failure")

        monkeypatch.setattr("src.pipeline.phase1_pipeline.validate_document", _boom)

        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        assert result.success is False
        assert result.failed_stage == "run_validation"
        assert result.markdown_path.is_file()
        assert result.docx_path.is_file()
        assert result.report_path is None


class TestOutputFileCreation:
    def test_markdown_file_is_written_under_markdown_subdir(self, tmp_path: Path) -> None:
        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        assert result.markdown_path == tmp_path / "markdown" / f"{A_SAMPLE_PDF.stem}.md"
        assert result.markdown_path.is_file()

    def test_docx_file_is_written_under_docx_subdir(self, tmp_path: Path) -> None:
        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        assert result.docx_path == tmp_path / "docx" / f"{A_SAMPLE_PDF.stem}.docx"
        assert result.docx_path.is_file()

    def test_report_file_is_written_under_reports_subdir(self, tmp_path: Path) -> None:
        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        assert result.report_path == tmp_path / "reports" / f"{A_SAMPLE_PDF.stem}.json"
        assert result.report_path.is_file()

    def test_images_are_written_under_images_subdir(self, tmp_path: Path) -> None:
        run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        image_dir = tmp_path / "images" / A_SAMPLE_PDF.stem
        assert image_dir.is_dir()
        assert any(image_dir.iterdir())


class TestMetadataSynchronization:
    def test_page_count_matches_actual_pages(self, tmp_path: Path) -> None:
        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        document = result.document
        assert document.metadata.page_count == len(document.pages)

    def test_image_count_matches_actual_images(self, tmp_path: Path) -> None:
        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        document = result.document
        assert document.metadata.image_count == len(document.images)
        assert document.metadata.image_count > 0  # this sample PDF has images


class TestValidationReportGeneration:
    def test_report_is_valid_json_with_expected_shape(self, tmp_path: Path) -> None:
        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        report = json.loads(result.report_path.read_text(encoding="utf-8"))

        assert report["source_pdf_path"] == str(A_SAMPLE_PDF)
        assert "generated_at" in report
        assert set(report["summary"].keys()) == {"error", "warning", "info", "total"}
        assert report["summary"]["total"] == len(result.validation_issues)
        assert len(report["issues"]) == len(result.validation_issues)

    def test_document_validation_issues_field_is_populated(self, tmp_path: Path) -> None:
        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        assert result.document.validation_issues == result.validation_issues

    def test_known_pipeline_gaps_are_flagged_for_real_pdfs(self, tmp_path: Path) -> None:
        # OCR is not implemented yet, so this real sample has images
        # (it extracts successfully, since this PDF is a full-page
        # scan) but no extracted text or content headings - the
        # validator should still run successfully and surface the
        # missing-H1 gap. DOC_001 ("empty document") correctly does
        # NOT fire here, since the document does have images.
        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        rule_ids = {issue.rule_id for issue in result.validation_issues}
        assert "HEADING_002" in rule_ids
        assert "DOC_001" not in rule_ids


class TestPhaseD0RoutingIntegration:
    """Confirms the OCR routing layer (Phase D.0) is actually wired into
    the live pipeline, not just unit-tested in isolation.
    """

    def test_scanned_pdf_routes_every_page_to_ocr_required(self, tmp_path: Path) -> None:
        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        assert all(page.page_type == PageType.OCR_REQUIRED for page in result.document.pages)

    def test_born_digital_pdf_routes_pages_to_direct_text(self, tmp_path: Path) -> None:
        result = run_pipeline(A_DIGITAL_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        assert all(page.page_type == PageType.DIRECT_TEXT for page in result.document.pages)

    def test_every_page_has_routing_metadata_populated(self, tmp_path: Path) -> None:
        result = run_pipeline(A_DIGITAL_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)
        for page in result.document.pages:
            assert page.page_type is not None
            assert page.extraction_method is not None
            assert page.routing_decision is not None


@pytest.mark.real_docling
class TestDoclingIntegration:
    """Real, unmocked Docling calls through the full live pipeline
    (Phase D.1) - not just src/ocr/docling_engine.py in isolation.
    Scoped to a single extracted page to keep runtime bounded (see
    tests/test_docling_engine.py's TestRealDoclingIntegration for why).
    """

    @pytest.fixture
    def single_scanned_page_pdf(self, tmp_path: Path) -> Path:
        import fitz

        single_page_pdf = tmp_path / "oleary_page2.pdf"
        with fitz.open(A_SAMPLE_PDF) as src:
            single_doc = fitz.open()
            single_doc.insert_pdf(src, from_page=1, to_page=1)  # 0-indexed: page 2
            single_doc.save(str(single_page_pdf))
            single_doc.close()
        return single_page_pdf

    def test_enable_ocr_true_recovers_real_text_through_full_pipeline(
        self, single_scanned_page_pdf: Path, tmp_path: Path
    ) -> None:
        result = run_pipeline(single_scanned_page_pdf, output_root=tmp_path, enable_ocr=True)

        assert result.success is True
        page = result.document.pages[0]
        assert page.page_type == PageType.OCR_REQUIRED
        assert len(page.cleaned_text.strip()) > 0
        assert result.ocr_metrics is not None
        assert result.ocr_metrics.page_count == 1

    def test_enable_ocr_false_leaves_text_empty(
        self, single_scanned_page_pdf: Path, tmp_path: Path
    ) -> None:
        result = run_pipeline(single_scanned_page_pdf, output_root=tmp_path, enable_ocr=False)

        assert result.success is True
        page = result.document.pages[0]
        assert page.cleaned_text == ""
        assert result.ocr_metrics is None


class TestSuryaFallbackWiring:
    """Confirms the Surya fallback (Phase D.2) is actually wired into the
    live pipeline, not just unit-tested in isolation. Mocks
    run_docling_ocr/run_surya_ocr at the phase1_pipeline module
    boundary (matching TestStageFailureHandling's pattern above) so
    these run fast and deterministically, without needing real OCR.
    """

    def test_surya_runs_after_docling_when_ocr_enabled(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        call_order = []

        def _fake_docling(document, metrics=None):
            call_order.append("docling")
            return document

        def _fake_surya(document, metrics=None):
            call_order.append("surya")
            return document

        monkeypatch.setattr("src.pipeline.phase1_pipeline.run_docling_ocr", _fake_docling)
        monkeypatch.setattr("src.pipeline.phase1_pipeline.run_surya_ocr", _fake_surya)

        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=True)

        assert call_order == ["docling", "surya"]
        assert result.success is True
        assert result.ocr_metrics is not None
        assert result.surya_metrics is not None

    def test_surya_does_not_run_when_ocr_disabled(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        surya_called = {"n": 0}

        def _fail_if_called(document, metrics=None):
            surya_called["n"] += 1
            raise AssertionError("run_surya_ocr should never be called when enable_ocr=False")

        monkeypatch.setattr("src.pipeline.phase1_pipeline.run_surya_ocr", _fail_if_called)

        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        assert result.success is True
        assert surya_called["n"] == 0
        assert result.surya_metrics is None

    def test_surya_failure_is_caught_and_reported_at_extract_text_stage(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        def _fake_docling(document, metrics=None):
            return document

        def _boom(document, metrics=None):
            raise RuntimeError("simulated Surya fallback failure")

        monkeypatch.setattr("src.pipeline.phase1_pipeline.run_docling_ocr", _fake_docling)
        monkeypatch.setattr("src.pipeline.phase1_pipeline.run_surya_ocr", _boom)

        result = run_pipeline(A_SAMPLE_PDF, output_root=tmp_path, enable_ocr=True)

        assert result.success is False
        assert result.failed_stage == "extract_text"
        assert "simulated Surya fallback failure" in result.error_message
        # Docling's own metrics survive even though the later Surya call failed
        assert result.ocr_metrics is not None


@pytest.mark.real_surya
class TestSuryaIntegration:
    """Real, unmocked Surya calls through the full live pipeline (Phase
    D.2) - not just src/ocr/surya_engine.py in isolation. Docling itself
    is mocked here to deterministically force its "attempted but empty"
    state on a real benchmark page (rather than depending on whether
    real Docling happens to fail on a particular page), so this test
    exercises real Surya inference against genuine benchmark content
    while keeping the trigger condition reliable. Scoped to a single
    extracted page to keep runtime bounded (see
    tests/test_surya_engine.py's TestRealSuryaIntegration for why).
    """

    @pytest.fixture
    def single_scanned_page_pdf(self, tmp_path: Path) -> Path:
        import fitz

        single_page_pdf = tmp_path / "oleary_page2.pdf"
        with fitz.open(A_SAMPLE_PDF) as src:
            single_doc = fitz.open()
            single_doc.insert_pdf(src, from_page=1, to_page=1)  # 0-indexed: page 2
            single_doc.save(str(single_page_pdf))
            single_doc.close()
        return single_page_pdf

    def test_docling_empty_falls_back_to_real_surya_through_full_pipeline(
        self, single_scanned_page_pdf: Path, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        def _fake_docling_leaves_page_empty(document, metrics=None):
            for page in document.pages:
                page.extraction_method = ExtractionMethod.DOCLING
            return document

        monkeypatch.setattr(
            "src.pipeline.phase1_pipeline.run_docling_ocr", _fake_docling_leaves_page_empty
        )

        result = run_pipeline(single_scanned_page_pdf, output_root=tmp_path, enable_ocr=True)

        assert result.success is True
        page = result.document.pages[0]
        assert page.page_type == PageType.OCR_REQUIRED
        assert len(page.cleaned_text.strip()) > 0
        assert page.ocr_confidence == OCRConfidence.LOW
        assert page.extraction_method == ExtractionMethod.SURYA
        assert result.surya_metrics is not None
        assert result.surya_metrics.page_count == 1


class TestStructureDetectionWiring:
    """Confirms the Structure Detection stage (Phase H) is actually
    wired into the live pipeline, not just unit-tested in isolation.
    Mocks src.pipeline.phase1_pipeline.detect_structure at the module
    boundary (matching TestStageFailureHandling's pattern above) so
    these run fast and deterministically.
    """

    def test_detect_structure_runs_between_ocr_and_image_extraction(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        call_order = []

        def _fake_structure(document):
            call_order.append("detect_structure")
            return document

        def _fake_images(document, output_dir):
            call_order.append("extract_images")
            return document

        monkeypatch.setattr("src.pipeline.phase1_pipeline.detect_structure", _fake_structure)
        monkeypatch.setattr("src.pipeline.phase1_pipeline.extract_images", _fake_images)

        result = run_pipeline(A_DIGITAL_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        assert call_order == ["detect_structure", "extract_images"]
        assert result.success is True

    def test_document_blocks_populated_for_real_benchmark_pdf(self, tmp_path: Path) -> None:
        result = run_pipeline(A_DIGITAL_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        assert result.success is True
        assert len(result.document.blocks) > 0
        assert all(block.page_number >= 1 for block in result.document.blocks)

    def test_structure_detection_failure_is_caught_and_reported(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        def _boom(document):
            raise RuntimeError("simulated structure detection failure")

        monkeypatch.setattr("src.pipeline.phase1_pipeline.detect_structure", _boom)

        result = run_pipeline(A_DIGITAL_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        assert result.success is False
        assert result.failed_stage == "detect_structure"
        assert "simulated structure detection failure" in result.error_message
        # no later stage ran
        assert result.markdown_path is None
        assert result.docx_path is None


@pytest.mark.parametrize("sample_pdf_path", SAMPLE_PDFS, ids=[p.name for p in SAMPLE_PDFS])
class TestStructureDetectionDoesNotChangeExistingOutputs:
    """The core Phase H regression guard: this stage is purely additive
    metadata, so disabling it (forcing it to a no-op passthrough,
    reproducing the exact pre-Phase-H pipeline shape) must produce
    byte-identical markdown, identical headings, identical validation
    issues, and the same image count as running it for real.
    """

    def test_markdown_headings_and_validation_are_unaffected(
        self, sample_pdf_path: Path, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        # Same output_root for both runs: image extraction is
        # deterministic, so re-running it into the same directory
        # reproduces identical absolute image paths in markdown,
        # keeping this a clean comparison of Structure Detection's
        # effect only - not an artifact of two different tmp_path trees.
        result_with = run_pipeline(sample_pdf_path, output_root=tmp_path, enable_ocr=False)

        monkeypatch.setattr(
            "src.pipeline.phase1_pipeline.detect_structure", lambda document: document
        )
        result_without = run_pipeline(sample_pdf_path, output_root=tmp_path, enable_ocr=False)

        assert result_with.success is True
        assert result_without.success is True

        assert result_with.markdown_path.read_text(
            encoding="utf-8"
        ) == result_without.markdown_path.read_text(encoding="utf-8")

        # feature_009: an H6 page marker's text legitimately differs
        # between these two runs whenever Structure Detection populated
        # Page.printed_label - "with" may show the real printed page
        # number (e.g. "343"), "without" always falls back to the
        # physical page_number (printed_label is never populated when
        # detect_structure() is stubbed to a no-op). This is the
        # documented, intended behavior (see
        # samples/regressions/feature_009_printed_page_number_preservation/
        # notes_md/printed_page_number_audit.md), not Structure
        # Detection changing existing output it shouldn't - excluded
        # from this comparison the same documented-exception way
        # PAGE_003/DOC_004/NOTE_001/NOTE_002 are excluded from the
        # validation-issue comparison below. Every other heading field,
        # and every non-page-marker heading's text, is still compared
        # exactly - this is the actual regression guard's target.
        def _heading_key(heading):
            data = heading.model_dump()
            if data["is_page_marker"]:
                data = {**data, "text": None}
            return data

        assert [_heading_key(h) for h in result_with.document.headings] == [
            _heading_key(h) for h in result_without.document.headings
        ]
        # Compared without "message": IMAGE_xxx messages embed each
        # image's randomly-generated image_id (see image_extractor.py),
        # which legitimately differs between these two separate
        # run_pipeline() calls regardless of Structure Detection -
        # unrelated to what this test is checking. severity/rule_id/
        # page_number/suggested_action fully capture "the same issues
        # fired the same way".
        def _issue_key(issue):
            return (issue.severity, issue.rule_id, issue.page_number, issue.suggested_action)

        # PAGE_003 (Phase I.1, reading-order anomalies), NOTE_001/NOTE_002
        # (Phase K, footnote/endnote detection), and DOC_004 (XML
        # Sanitization Architecture, Layer 2) are excluded from this
        # comparison on purpose: unlike every other rule, all are real,
        # deliberate consumers/side-effects of Structure Detection
        # actually running. PAGE_003 reads document.blocks directly.
        # NOTE_001/002 come from footnote_detector.py, which reads
        # document.blocks and explicitly no-ops ("No structure blocks
        # available... skipping footnote detection") when it's empty -
        # exactly what the "without" stub guarantees. DOC_004 discloses
        # sanitization events that structure_detector.py's own text
        # re-derivation pass is one of several independent producers of
        # (see src/utils/text_sanitization.py's Layer 1 call sites). All
        # are exactly the output this test's "without" run forces to
        # stay empty/never-run. So a multi-column PDF, a PDF with real
        # footnotes/endnotes, or a PDF whose structure-detection
        # re-derivation pass independently hits an XML-illegal
        # character, genuinely produces issues in result_with that
        # result_without can never produce; that divergence is those
        # rules working as designed, not Structure Detection "changing
        # existing outputs" it shouldn't. Confirmed via the benchmark
        # corpus once it grew to include real multi-column and
        # footnote-bearing PDFs (Nature of Enquiry, Aims of Education,
        # sockett, Brinkman) - the original 4-PDF set never exposed any
        # of these three cases.
        #
        # HEADING_004 (feature_009) joins this set for the same reason:
        # _check_duplicate_headings() is explicitly scoped to include
        # H6 page markers, and printed_label (only ever populated when
        # Structure Detection actually runs) lets two different
        # physical pages legitimately share identical marker text (e.g.
        # sockett_profession.pdf's physical pages 3 and 5 both print
        # the roman numeral "I") - something that could never happen
        # under plain physical page_number, which is always unique per
        # page. "with" correctly flags this real duplicate; "without"
        # can never produce it since printed_label is never populated.
        _EXCLUDED_RULE_IDS = {"PAGE_003", "DOC_004", "NOTE_001", "NOTE_002", "HEADING_004"}
        with_keys = [
            _issue_key(i) for i in result_with.validation_issues if i.rule_id not in _EXCLUDED_RULE_IDS
        ]
        without_keys = [
            _issue_key(i)
            for i in result_without.validation_issues
            if i.rule_id not in _EXCLUDED_RULE_IDS
        ]
        assert with_keys == without_keys
        # PAGE_003 and NOTE_001/NOTE_002 specifically can never fire in
        # "without": both read document.blocks (directly, or via
        # footnote_detector.py's explicit empty-blocks no-op), which the
        # stub guarantees stays empty. DOC_004 has no equivalent
        # guarantee - extraction (extractor.py/docling_engine.py/
        # surya_engine.py) still runs in both branches and is itself an
        # independent Layer 1 sanitizer, so DOC_004 may legitimately
        # fire in "without" too; only excluded above from the equality
        # check, not asserted absent.
        assert all(
            i.rule_id not in {"PAGE_003", "NOTE_001", "NOTE_002"}
            for i in result_without.validation_issues
        )
        assert len(result_with.document.images) == len(result_without.document.images)

        # The only expected difference: Structure Detection actually ran
        # for "with" (block counts covered separately by
        # tests/test_structure_detector.py and
        # TestStructureDetectionWiring above) and was forced to a no-op
        # for "without".
        assert result_without.document.blocks == []


class TestAltTextDatasetSidecar:
    """Phase F.5: a filesystem dataset sidecar is written as a
    side-output of the Extract Images stage."""

    def test_dataset_path_populated_on_success(self, tmp_path: Path) -> None:
        result = run_pipeline(A_DIGITAL_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        assert result.success is True
        assert result.alt_text_dataset_path is not None
        assert result.alt_text_dataset_path.is_file()

    def test_dataset_path_is_none_when_extract_images_fails(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        def _boom(document, output_dir):
            raise RuntimeError("simulated image extraction failure")

        monkeypatch.setattr("src.pipeline.phase1_pipeline.extract_images", _boom)

        result = run_pipeline(A_DIGITAL_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        assert result.success is False
        assert result.alt_text_dataset_path is None

    def test_dataset_json_shape_matches_real_extracted_images(self, tmp_path: Path) -> None:
        result = run_pipeline(A_DIGITAL_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        dataset = json.loads(result.alt_text_dataset_path.read_text(encoding="utf-8"))
        assert dataset["source_pdf_path"] == str(A_DIGITAL_SAMPLE_PDF)
        assert "generated_at" in dataset
        assert len(dataset["images"]) == len(result.document.images)

        for record in dataset["images"]:
            assert record["figure"]["alt_text_status"] == "pending_review"
            assert record["figure"]["alt_text"]
            assert "nearby_text" in record

    def test_dataset_written_even_for_documents_with_no_images(self, tmp_path: Path) -> None:
        no_images_pdf = SAMPLE_PDF_DIR / "5.Teachingas a profession_Calderhead.pdf"
        result = run_pipeline(no_images_pdf, output_root=tmp_path, enable_ocr=False)

        assert result.document.images == []
        dataset = json.loads(result.alt_text_dataset_path.read_text(encoding="utf-8"))
        assert dataset["images"] == []

    def test_failed_extractions_are_excluded_from_dataset(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        import fitz

        image_bearing_pdf = SAMPLE_PDF_DIR / "4.Teaching as a professional discipline-Chapter 1.pdf"

        def _always_fail(self, xref):
            raise RuntimeError("simulated corrupt image data")

        monkeypatch.setattr(fitz.Document, "extract_image", _always_fail)

        result = run_pipeline(image_bearing_pdf, output_root=tmp_path, enable_ocr=False)

        assert any(image.extraction_failed for image in result.document.images)
        dataset = json.loads(result.alt_text_dataset_path.read_text(encoding="utf-8"))
        assert dataset["images"] == []


@pytest.mark.parametrize("sample_pdf_path", SAMPLE_PDFS, ids=[p.name for p in SAMPLE_PDFS])
class TestAltTextDoesNotChangeExistingOutputsOutsideImages:
    """The core Phase F.1-F.5 regression guard: figure/caption/alt-text
    linking is purely additive to image accessibility metadata. Forcing
    it to a no-op (reproducing the exact pre-Phase-F image.figure=None
    behavior) must produce byte-identical non-image markdown lines and
    identical non-IMAGE_004 validation issues.
    """

    def test_non_image_outputs_are_unaffected(
        self, sample_pdf_path: Path, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        result_with = run_pipeline(sample_pdf_path, output_root=tmp_path, enable_ocr=False)

        def _no_op_link_figures(images, blocks_by_page):
            return None  # every image.figure stays None, exactly pre-Phase-F behavior

        monkeypatch.setattr("src.images.image_extractor._link_figures", _no_op_link_figures)
        result_without = run_pipeline(sample_pdf_path, output_root=tmp_path, enable_ocr=False)

        assert result_with.success is True
        assert result_without.success is True

        lines_with = result_with.markdown_path.read_text(encoding="utf-8").splitlines()
        lines_without = result_without.markdown_path.read_text(encoding="utf-8").splitlines()
        non_image_with = [line for line in lines_with if not line.startswith("![")]
        non_image_without = [line for line in lines_without if not line.startswith("![")]
        assert non_image_with == non_image_without

        # Image lines themselves must still reference the same files,
        # in the same order - only the alt-text bracket content differs.
        image_paths_with = [line.split("](", 1)[1] for line in lines_with if line.startswith("![")]
        image_paths_without = [
            line.split("](", 1)[1] for line in lines_without if line.startswith("![")
        ]
        assert image_paths_with == image_paths_without

        issues_with = {(i.rule_id, i.message, i.page_number) for i in result_with.validation_issues if i.rule_id != "IMAGE_004"}
        issues_without = {
            (i.rule_id, i.message, i.page_number)
            for i in result_without.validation_issues
            if i.rule_id != "IMAGE_004"
        }
        assert issues_with == issues_without

        # IMAGE_004 only ever appears once Figure/alt-text linking
        # actually runs - confirms the new check is additive, not a
        # silent behavior change to anything pre-existing.
        pending_with = [i for i in result_with.validation_issues if i.rule_id == "IMAGE_004"]
        pending_without = [i for i in result_without.validation_issues if i.rule_id == "IMAGE_004"]
        assert len(pending_with) == len(result_with.document.images)
        assert pending_without == []


class TestFootnoteDetectionWiring:
    """Confirms Footnote/Endnote Detection (Phase K) is actually wired
    into the live pipeline (inside the existing Detect Structure stage,
    not a new one), not just unit-tested in isolation."""

    def test_detect_footnotes_runs_after_detect_structure(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        call_order = []

        def _fake_structure(document):
            call_order.append("detect_structure")
            return document

        def _fake_footnotes(document):
            call_order.append("detect_footnotes")
            return document

        monkeypatch.setattr("src.pipeline.phase1_pipeline.detect_structure", _fake_structure)
        monkeypatch.setattr("src.pipeline.phase1_pipeline.detect_footnotes", _fake_footnotes)

        result = run_pipeline(A_DIGITAL_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        assert call_order == ["detect_structure", "detect_footnotes"]
        assert result.success is True

    def test_footnote_detection_failure_is_caught_under_detect_structure_stage(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        def _boom(document):
            raise RuntimeError("simulated footnote detection failure")

        monkeypatch.setattr("src.pipeline.phase1_pipeline.detect_footnotes", _boom)

        result = run_pipeline(A_DIGITAL_SAMPLE_PDF, output_root=tmp_path, enable_ocr=False)

        assert result.success is False
        # Bundled into the same stage as Structure Detection - no new
        # pipeline stage was added for this phase.
        assert result.failed_stage == "detect_structure"
        assert "simulated footnote detection failure" in result.error_message


def test_real_footnote_is_detected_and_rendered_end_to_end(tmp_path: Path) -> None:
    # The exact worked example from the Phase K brief, run through the
    # complete, real pipeline (parse -> structure -> footnotes ->
    # images -> headings -> markdown -> docx -> validation).
    import fitz

    superscript_one = chr(0xB9)  # "¹" SUPERSCRIPT ONE
    pdf_path = tmp_path / "footnote.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Doc Title", fontname="helv", fontsize=14)
    # Two extra 12pt filler lines: src/footnotes/footnote_detector.py
    # picks the document's dominant body font size by a character-count-
    # weighted vote (mirroring heading_detector.py's body-profile vote) -
    # with only one 12pt line, the 8pt footnote line below would have
    # *more* characters and incorrectly win the vote as "the body font".
    page.insert_text((72, 92), "This is the first line of ordinary body text.", fontname="helv", fontsize=12)
    page.insert_text((72, 112), "This is the second line of ordinary body text.", fontname="helv", fontsize=12)
    page.insert_text(
        (72, 132),
        f"The study showed significant improvement{superscript_one}.",
        fontname="helv",
        fontsize=12,
    )
    page.insert_text(
        (72, 700),
        f"{superscript_one} Improvement measured using standardized test scores.",
        fontname="helv",
        fontsize=8,
    )
    doc.save(str(pdf_path))
    doc.close()

    result = run_pipeline(pdf_path, output_root=tmp_path, enable_ocr=False)

    assert result.success is True
    assert len(result.document.footnotes) == 1
    note = result.document.footnotes[0]
    assert note.body == "Improvement measured using standardized test scores."

    markdown = result.markdown_path.read_text(encoding="utf-8")
    assert "improvement[^p1-1]." in markdown
    assert "[^p1-1]: Improvement measured using standardized test scores." in markdown

    assert any(issue.rule_id == "NOTE_001" for issue in result.validation_issues)

    import zipfile
    from lxml import etree as _etree

    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    import docx as _docx
    docx_doc = _docx.Document(str(result.docx_path))

    # Native footnote reference must exist in the body; no legacy hyperlinks.
    refs = [
        el
        for paragraph in docx_doc.paragraphs
        for el in paragraph._p.findall(f".//{w_ns}footnoteReference")
    ]
    assert len(refs) == 1

    hyperlinks = [
        el
        for paragraph in docx_doc.paragraphs
        for el in paragraph._p.findall(f".//{w_ns}hyperlink")
    ]
    assert len(hyperlinks) == 0

    # footnotes.xml part must be present and contain the body text.
    with zipfile.ZipFile(str(result.docx_path)) as zf:
        assert "word/footnotes.xml" in zf.namelist()
        root = _etree.fromstring(zf.read("word/footnotes.xml"))
    all_text = "".join(t.text or "" for t in root.iter(f"{w_ns}t"))
    assert "Improvement measured using standardized test scores." in all_text


@pytest.mark.parametrize("sample_pdf_path", SAMPLE_PDFS, ids=[p.name for p in SAMPLE_PDFS])
class TestFootnoteDetectionBenchmarkRegression:
    """Confirmed during the Phase K architecture audit: none of the
    (then 4) benchmark PDFs contain a real footnote or endnote. Pinned
    down here as a regression guard against the live pipeline, not just
    the detector in isolation. The benchmark corpus grew to 10 PDFs on
    2026-06-24 (see DECISIONS_LOG.md "Benchmark Corpus Expansion");
    Brinkman - now also in this corpus, not just samples/regressions/ -
    has 3 real, confirmed, body-linked endnotes (the exact bug_005 fix
    this detector exists for), so it's excluded below, not because
    detection is wrong, but because it would be wrong for this specific
    PDF to report zero."""

    _PDF_WITH_REAL_FOOTNOTES = (
        "7.brinkman-learner-centred-education-reform-india-missing-beliefs.pdf"
    )

    def test_no_footnotes_detected_and_no_note_issues_raised(
        self, sample_pdf_path: Path, tmp_path: Path
    ) -> None:
        if sample_pdf_path.name == self._PDF_WITH_REAL_FOOTNOTES:
            pytest.skip("known to have real, body-linked footnotes/endnotes - see bug_005")
        result = run_pipeline(sample_pdf_path, output_root=tmp_path, enable_ocr=False)

        assert result.success is True
        assert result.document.footnotes == []
        assert [i for i in result.validation_issues if i.rule_id.startswith("NOTE_")] == []
