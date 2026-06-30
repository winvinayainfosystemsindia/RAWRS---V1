"""Tests for FEATURE_015.1 — Semantic Accessible Table Remediation.

Covers:
  - TABLE_001–005 validation rules (Phase C)
  - Markdown table-id comment emission (Phase A / Task 5)
  - DOCX w:tblHeader on header rows (Phase F)
  - DOCX caption and summary paragraphs (Phase F)
  - DOCX cell merges via col_span / row_span (Phase A)
  - AI analyzer stub output (Phase B)
  - Screen reader announcement logic (Phase E, Python-model side)
  - API analyze endpoint (Phase B / Task 7)
  - API header_col_count via PATCH (Task 7)
"""

import io
import os
import zipfile
from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from fastapi.testclient import TestClient

from src.api.jobs import _jobs, Job, JobStatus
from src.api.main import app
from src.models.contracts import Document, Metadata, ProcessingStatus, Severity
from src.models.table import Table, TableAISuggestions, TableCell, TableRow, TableStatus
from src.pipeline.phase1_pipeline import PipelineResult
from src.validation.validator import validate_document


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture(autouse=True)
def clear_jobs():
    _jobs.clear()
    yield
    _jobs.clear()


@pytest.fixture()
def client():
    return TestClient(app)


def _make_minimal_table(
    table_id: str = "table-p1-0",
    has_caption: bool = False,
    has_summary: bool = False,
    has_header_row: bool = False,
    confidence: float = 1.0,
    status: TableStatus = TableStatus.AUTO_DETECTED,
) -> Table:
    cells_h = [
        TableCell(text="Col A", row_index=0, col_index=0, is_header=True),
        TableCell(text="Col B", row_index=0, col_index=1, is_header=True),
    ]
    cells_d = [
        TableCell(text="1", row_index=1, col_index=0),
        TableCell(text="2", row_index=1, col_index=1),
    ]
    return Table(
        table_id=table_id,
        page_number=1,
        row_count=2,
        col_count=2,
        rows=[
            TableRow(cells=cells_h, is_header_row=has_header_row),
            TableRow(cells=cells_d, is_header_row=False),
        ],
        caption="My caption" if has_caption else None,
        summary="My summary" if has_summary else None,
        confidence=confidence,
        status=status,
    )


def _make_doc(table: Table) -> Document:
    doc = Document(
        source_pdf_path="test.pdf",
        metadata=Metadata(filename="test.pdf", page_count=1),
    )
    doc.tables = [table]
    return doc


def _make_job_with_table(table: Table) -> str:
    doc = _make_doc(table)
    job = Job(
        job_id="job-test-001",
        filename="test.pdf",
        pdf_path=Path("test.pdf"),
        status=JobStatus.COMPLETE,
        created_at=datetime.now(timezone.utc),
        result=PipelineResult(
            source_pdf_path="test.pdf",
            success=True,
            status=ProcessingStatus.VALIDATED,
            duration_seconds=0.1,
            document=doc,
            markdown_path=None,
            docx_path=None,
            report_path=None,
            alt_text_dataset_path=None,
        ),
    )
    _jobs["job-test-001"] = job
    return "job-test-001"


# ---------------------------------------------------------------------------
# Phase C: Validation rules TABLE_001–005
# ---------------------------------------------------------------------------


class TestTableValidationRules:
    def _issues(self, table: Table) -> list:
        doc = _make_doc(table)
        return validate_document(doc)

    def test_TABLE_001_fires_when_no_caption(self):
        table = _make_minimal_table(has_caption=False, has_header_row=True, has_summary=True)
        issues = self._issues(table)
        rule_ids = [i.rule_id for i in issues]
        assert "TABLE_001" in rule_ids
        match = next(i for i in issues if i.rule_id == "TABLE_001")
        assert match.severity == Severity.WARNING
        assert "caption" in match.message.lower()

    def test_TABLE_001_absent_when_caption_present(self):
        table = _make_minimal_table(has_caption=True, has_header_row=True, has_summary=True)
        issues = self._issues(table)
        assert not any(i.rule_id == "TABLE_001" for i in issues)

    def test_TABLE_002_fires_when_no_summary(self):
        table = _make_minimal_table(has_caption=True, has_header_row=True, has_summary=False)
        issues = self._issues(table)
        rule_ids = [i.rule_id for i in issues]
        assert "TABLE_002" in rule_ids
        match = next(i for i in issues if i.rule_id == "TABLE_002")
        assert match.severity == Severity.WARNING
        assert "summary" in match.message.lower() or "h73" in match.message.lower()

    def test_TABLE_002_absent_when_summary_present(self):
        table = _make_minimal_table(has_caption=True, has_header_row=True, has_summary=True)
        issues = self._issues(table)
        assert not any(i.rule_id == "TABLE_002" for i in issues)

    def test_TABLE_003_fires_when_no_header_row(self):
        table = _make_minimal_table(has_caption=True, has_summary=True, has_header_row=False)
        issues = self._issues(table)
        rule_ids = [i.rule_id for i in issues]
        assert "TABLE_003" in rule_ids
        match = next(i for i in issues if i.rule_id == "TABLE_003")
        assert match.severity == Severity.WARNING

    def test_TABLE_003_absent_when_header_row_present(self):
        table = _make_minimal_table(has_caption=True, has_summary=True, has_header_row=True)
        issues = self._issues(table)
        assert not any(i.rule_id == "TABLE_003" for i in issues)

    def test_TABLE_004_fires_for_empty_header_cell(self):
        empty_header_cells = [
            TableCell(text="", row_index=0, col_index=0, is_header=True),
            TableCell(text="Value", row_index=0, col_index=1, is_header=True),
        ]
        data_cells = [
            TableCell(text="1", row_index=1, col_index=0),
            TableCell(text="2", row_index=1, col_index=1),
        ]
        table = Table(
            table_id="table-p1-0",
            page_number=1,
            row_count=2,
            col_count=2,
            rows=[
                TableRow(cells=empty_header_cells, is_header_row=True),
                TableRow(cells=data_cells, is_header_row=False),
            ],
            caption="Cap",
            summary="Sum",
            confidence=1.0,
        )
        issues = self._issues(table)
        rule_ids = [i.rule_id for i in issues]
        assert "TABLE_004" in rule_ids
        match = next(i for i in issues if i.rule_id == "TABLE_004")
        assert match.severity == Severity.WARNING
        assert "empty" in match.message.lower()

    def test_TABLE_004_absent_when_all_header_cells_filled(self):
        table = _make_minimal_table(has_caption=True, has_summary=True, has_header_row=True)
        issues = self._issues(table)
        assert not any(i.rule_id == "TABLE_004" for i in issues)

    def test_TABLE_005_fires_for_low_confidence_auto_detected(self):
        table = _make_minimal_table(
            has_caption=True, has_summary=True, has_header_row=True,
            confidence=0.5, status=TableStatus.AUTO_DETECTED
        )
        issues = self._issues(table)
        rule_ids = [i.rule_id for i in issues]
        assert "TABLE_005" in rule_ids
        match = next(i for i in issues if i.rule_id == "TABLE_005")
        assert match.severity == Severity.INFO
        assert "50%" in match.message or "confidence" in match.message.lower()

    def test_TABLE_005_absent_when_confidence_high(self):
        table = _make_minimal_table(
            has_caption=True, has_summary=True, has_header_row=True,
            confidence=0.9
        )
        issues = self._issues(table)
        assert not any(i.rule_id == "TABLE_005" for i in issues)

    def test_TABLE_005_absent_for_manually_created_even_if_confidence_zero(self):
        table = _make_minimal_table(
            has_caption=True, has_summary=True, has_header_row=True,
            confidence=0.0, status=TableStatus.MANUALLY_CREATED
        )
        issues = self._issues(table)
        assert not any(i.rule_id == "TABLE_005" for i in issues)

    def test_all_five_rules_fire_for_worst_case_table(self):
        empty_header = [
            TableCell(text="", row_index=0, col_index=0, is_header=True),
        ]
        data = [TableCell(text="x", row_index=1, col_index=0)]
        table = Table(
            table_id="table-p1-0",
            page_number=1,
            row_count=2,
            col_count=1,
            rows=[
                TableRow(cells=empty_header, is_header_row=True),
                TableRow(cells=data, is_header_row=False),
            ],
            caption=None,
            summary=None,
            confidence=0.3,
            status=TableStatus.AUTO_DETECTED,
        )
        issues = self._issues(table)
        rule_ids = {i.rule_id for i in issues}
        assert {"TABLE_001", "TABLE_002", "TABLE_004", "TABLE_005"} <= rule_ids
        # TABLE_003 absent because there IS a header row (even though a cell is empty)
        assert "TABLE_003" not in rule_ids

    def test_page_number_on_issue_matches_table_page(self):
        table = _make_minimal_table(has_caption=False, has_header_row=True, has_summary=True)
        table.page_number = 7
        issues = self._issues(table)
        t001 = next(i for i in issues if i.rule_id == "TABLE_001")
        assert t001.page_number == 7


# ---------------------------------------------------------------------------
# Phase A: Markdown table-id comment emission
# ---------------------------------------------------------------------------


class TestMarkdownTableId:
    def _markdown(self, table: Table) -> str:
        from src.markdown.markdown_builder import _render_tables
        blocks = _render_tables([table])
        return "\n".join(blocks)

    def test_table_id_comment_emitted_before_table(self):
        table = _make_minimal_table(has_header_row=True)
        table.table_id = "table-p3-1"
        md = self._markdown(table)
        lines = md.splitlines()
        id_comment_idx = next(
            (i for i, ln in enumerate(lines) if ln == "<!-- table-id: table-p3-1 -->"),
            None,
        )
        assert id_comment_idx is not None, "table-id comment missing"
        pipe_idx = next(i for i, ln in enumerate(lines) if ln.startswith("|"))
        assert id_comment_idx < pipe_idx, "table-id comment must come before pipe table"

    def test_caption_emitted_between_id_comment_and_pipe_table(self):
        table = _make_minimal_table(has_header_row=True, has_caption=True)
        table.table_id = "table-p1-0"
        md = self._markdown(table)
        lines = md.splitlines()
        id_idx = next(i for i, ln in enumerate(lines) if "table-id:" in ln)
        cap_idx = next((i for i, ln in enumerate(lines) if "*" in ln), None)
        pipe_idx = next(i for i, ln in enumerate(lines) if ln.startswith("|"))
        assert cap_idx is not None
        assert id_idx < cap_idx < pipe_idx

    def test_summary_emitted_as_comment_after_pipe_table(self):
        table = _make_minimal_table(has_header_row=True, has_summary=True)
        table.table_id = "table-p1-0"
        md = self._markdown(table)
        lines = md.splitlines()
        summary_idx = next(
            (i for i, ln in enumerate(lines) if "table-summary:" in ln),
            None,
        )
        assert summary_idx is not None
        last_pipe_idx = max(i for i, ln in enumerate(lines) if ln.startswith("|"))
        assert summary_idx > last_pipe_idx

    def test_empty_table_emits_nothing(self):
        table = Table(
            table_id="table-p1-0",
            page_number=1,
            row_count=0,
            col_count=0,
            rows=[],
        )
        from src.markdown.markdown_builder import _render_tables
        assert _render_tables([table]) == []


# ---------------------------------------------------------------------------
# Phase F: DOCX accessibility — w:tblHeader, caption, summary, merges
# ---------------------------------------------------------------------------


def _build_docx_xml(table: Table) -> str:
    """Generate a DOCX from a single table and return document.xml as a string."""
    from src.docx.docx_generator import _add_semantic_table

    docx_doc = DocxDocument()
    _add_semantic_table(docx_doc, table)

    buf = io.BytesIO()
    docx_doc.save(buf)
    buf.seek(0)
    with zipfile.ZipFile(buf) as zf:
        return zf.read("word/document.xml").decode("utf-8")


class TestDocxHeaderRow:
    def test_w_tblHeader_present_on_header_row(self):
        table = _make_minimal_table(has_header_row=True)
        xml = _build_docx_xml(table)
        assert "tblHeader" in xml, "w:tblHeader must be present for header rows"

    def test_w_tblHeader_absent_when_no_header_row(self):
        table = _make_minimal_table(has_header_row=False)
        xml = _build_docx_xml(table)
        assert "tblHeader" not in xml, "w:tblHeader must not appear when no header row"

    def test_header_row_cell_bold(self):
        table = _make_minimal_table(has_header_row=True)
        from src.docx.docx_generator import _add_semantic_table
        docx_doc = DocxDocument()
        _add_semantic_table(docx_doc, table)
        docx_table = docx_doc.tables[0]
        first_cell = docx_table.cell(0, 0)
        run = first_cell.paragraphs[0].runs[0]
        assert run.bold, "Header cell text must be bold"


class TestDocxCaption:
    def test_caption_paragraph_emitted_before_table(self):
        table = _make_minimal_table(has_header_row=True, has_caption=True)
        xml = _build_docx_xml(table)
        cap_pos = xml.find("My caption")
        tbl_pos = xml.find("<w:tbl>")
        assert cap_pos != -1, "Caption text not found in DOCX"
        assert cap_pos < tbl_pos, "Caption must appear before the table element"

    def test_no_caption_paragraph_when_caption_none(self):
        table = _make_minimal_table(has_header_row=True, has_caption=False)
        xml = _build_docx_xml(table)
        assert "My caption" not in xml

    def test_summary_paragraph_emitted_after_table(self):
        table = _make_minimal_table(has_header_row=True, has_summary=True)
        xml = _build_docx_xml(table)
        summary_pos = xml.find("My summary")
        tbl_end_pos = xml.find("</w:tbl>")
        assert summary_pos != -1, "Summary text not found in DOCX"
        assert summary_pos > tbl_end_pos, "Summary must appear after the table element"


class TestDocxMergedCells:
    def test_col_span_triggers_merge(self):
        cells = [
            TableCell(text="Merged", row_index=0, col_index=0, col_span=2, is_header=True),
            TableCell(text="", row_index=0, col_index=1, col_span=1, is_header=True),
        ]
        data = [
            TableCell(text="A", row_index=1, col_index=0),
            TableCell(text="B", row_index=1, col_index=1),
        ]
        table = Table(
            table_id="table-p1-0",
            page_number=1,
            row_count=2,
            col_count=2,
            rows=[
                TableRow(cells=cells, is_header_row=True),
                TableRow(cells=data, is_header_row=False),
            ],
        )
        # Should not raise and should produce a valid DOCX.
        xml = _build_docx_xml(table)
        assert "Merged" in xml

    def test_no_span_produces_unmerged_table(self):
        table = _make_minimal_table(has_header_row=True)
        from src.docx.docx_generator import _add_semantic_table
        docx_doc = DocxDocument()
        _add_semantic_table(docx_doc, table)
        assert len(docx_doc.tables) == 1
        assert len(docx_doc.tables[0].rows) == 2
        assert len(docx_doc.tables[0].rows[0].cells) == 2


# ---------------------------------------------------------------------------
# Phase F: DOCX via generate_docx (full pipeline, table-id comment tracking)
# ---------------------------------------------------------------------------


class TestDocxGeneratorTableId:
    def _generate(self, table: Table) -> str:
        """Build markdown → DOCX and return document.xml."""
        from src.markdown.markdown_builder import build_markdown
        from src.docx.docx_generator import generate_docx

        doc = _make_doc(table)
        from src.models.page import Page, PageType
        from src.models.contracts import ExtractionMethod
        doc.pages = [
            Page(
                page_number=1,
                raw_text="",
                page_type=PageType.DIRECT_TEXT,
                extraction_method=ExtractionMethod.DIRECT_TEXT_EXTRACTION,
            )
        ]
        from src.headings.heading_detector import detect_headings
        detect_headings(doc)
        md = build_markdown(doc)
        assert "<!-- table-id:" in md, "markdown must have table-id comment"

        buf_path = Path("C:/Users/WVF-D/AppData/Local/Temp/claude/test_table_docx.docx")
        buf_path.parent.mkdir(parents=True, exist_ok=True)
        generate_docx(doc, md, output_path=buf_path)

        buf = io.BytesIO(buf_path.read_bytes())
        with zipfile.ZipFile(buf) as zf:
            return zf.read("word/document.xml").decode("utf-8")

    def test_semantic_table_renders_via_generate_docx(self):
        table = _make_minimal_table(
            has_header_row=True, has_caption=True, has_summary=True
        )
        xml = self._generate(table)
        assert "tblHeader" in xml
        assert "My caption" in xml
        assert "My summary" in xml

    def test_caption_before_table_in_full_pipeline(self):
        table = _make_minimal_table(has_header_row=True, has_caption=True)
        xml = self._generate(table)
        cap_pos = xml.find("My caption")
        tbl_pos = xml.find("<w:tbl>")
        assert cap_pos < tbl_pos


# ---------------------------------------------------------------------------
# Phase B: AI analyzer stub
# ---------------------------------------------------------------------------


class TestAIAnalyzerStub:
    @pytest.fixture(autouse=True)
    def stub_env(self, monkeypatch):
        monkeypatch.setenv("RAWRS_AI_STUB", "1")

    def _request(self, cells, header_row_count=0, header_col_count=0, caption=None):
        from src.ai.table_analyzer import TableAnalysisRequest, analyze_table
        return analyze_table(TableAnalysisRequest(
            table_id="table-p1-0",
            page_number=1,
            row_count=len(cells),
            col_count=len(cells[0]) if cells else 0,
            header_row_count=header_row_count,
            header_col_count=header_col_count,
            cells=cells,
            existing_caption=caption,
            image_path=None,
        ))

    def test_stub_returns_result_without_model(self):
        result = self._request([["Name", "Value"], ["Alice", "42"]])
        assert result.table_type in {"simple", "complex", "data", "layout", "unknown"}
        assert isinstance(result.confidence, float)
        assert 0.0 <= result.confidence <= 1.0

    def test_numeric_body_detected_as_data_table(self):
        cells = [["Year", "Revenue"], ["2020", "100"], ["2021", "200"], ["2022", "350"]]
        result = self._request(cells, header_row_count=1)
        assert result.table_type == "data"

    def test_header_rows_passed_through(self):
        cells = [["A", "B"], ["1", "2"]]
        result = self._request(cells, header_row_count=1)
        assert result.header_rows_detected == 1

    def test_suggested_summary_always_present(self):
        result = self._request([["X"], ["y"]])
        assert result.suggested_summary is not None
        assert len(result.suggested_summary) > 0

    def test_no_caption_warning_when_caption_missing(self):
        result = self._request([["A", "B"], ["1", "2"]], caption=None)
        assert any("caption" in w.lower() for w in result.warnings)

    def test_no_caption_warning_absent_when_caption_provided(self):
        result = self._request([["A", "B"], ["1", "2"]], caption="My table")
        assert not any("caption" in w.lower() for w in result.warnings)

    def test_empty_table_returns_gracefully(self):
        result = self._request([])
        assert result.table_type == "unknown"
        assert result.header_rows_detected == 0


# ---------------------------------------------------------------------------
# Phase B / Task 7: API analyze endpoint
# ---------------------------------------------------------------------------


class TestAnalyzeTableEndpoint:
    @pytest.fixture(autouse=True)
    def stub_env(self, monkeypatch):
        monkeypatch.setenv("RAWRS_AI_STUB", "1")

    def test_analyze_returns_ai_suggestions(self, client):
        table = _make_minimal_table(has_header_row=True)
        job_id = _make_job_with_table(table)
        resp = client.post(f"/api/documents/{job_id}/tables/{table.table_id}/analyze")
        assert resp.status_code == 200
        data = resp.json()
        assert data["ai_suggestions"] is not None
        assert "table_type" in data["ai_suggestions"]
        assert "confidence" in data["ai_suggestions"]

    def test_analyze_404_on_unknown_table(self, client):
        table = _make_minimal_table()
        job_id = _make_job_with_table(table)
        resp = client.post(f"/api/documents/{job_id}/tables/nonexistent/analyze")
        assert resp.status_code == 404

    def test_analyze_422_on_empty_table(self, client):
        table = Table(
            table_id="table-p1-empty",
            page_number=1,
            row_count=0,
            col_count=0,
            rows=[],
        )
        job_id = _make_job_with_table(table)
        resp = client.post(f"/api/documents/{job_id}/tables/{table.table_id}/analyze")
        assert resp.status_code == 422

    def test_analyze_returns_updated_table_fields(self, client):
        table = _make_minimal_table(has_header_row=True)
        job_id = _make_job_with_table(table)
        resp = client.post(f"/api/documents/{job_id}/tables/{table.table_id}/analyze")
        data = resp.json()
        assert data["table_id"] == table.table_id
        assert data["row_count"] == 2
        assert data["col_count"] == 2


# ---------------------------------------------------------------------------
# Task 7: API header_col_count via PATCH
# ---------------------------------------------------------------------------


class TestHeaderColCountPatch:
    def test_patch_header_col_count_updates_model(self, client):
        table = _make_minimal_table(has_header_row=True)
        job_id = _make_job_with_table(table)
        resp = client.patch(
            f"/api/documents/{job_id}/tables/{table.table_id}",
            json={"header_col_count": 1},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["header_col_count"] == 1

    def test_patch_header_col_count_marks_first_col_as_row_header(self, client):
        table = _make_minimal_table(has_header_row=True)
        job_id = _make_job_with_table(table)
        client.patch(
            f"/api/documents/{job_id}/tables/{table.table_id}",
            json={"header_col_count": 1},
        )
        resp = client.get(f"/api/documents/{job_id}/tables")
        data = resp.json()["tables"][0]
        # Row 1 (data row, not header row) col 0 should be is_row_header=True
        data_row = next(r for r in data["rows"] if not r["is_header_row"])
        assert data_row["cells"][0]["is_row_header"] is True
        assert data_row["cells"][1]["is_row_header"] is False

    def test_patch_header_col_count_zero_clears_row_headers(self, client):
        table = _make_minimal_table(has_header_row=True)
        table.header_col_count = 1
        for row in table.rows:
            for cell in row.cells:
                if cell.col_index == 0 and not row.is_header_row:
                    cell.is_row_header = True
        job_id = _make_job_with_table(table)
        client.patch(
            f"/api/documents/{job_id}/tables/{table.table_id}",
            json={"header_col_count": 0},
        )
        resp = client.get(f"/api/documents/{job_id}/tables")
        data = resp.json()["tables"][0]
        for row in data["rows"]:
            for cell in row["cells"]:
                assert cell["is_row_header"] is False


# ---------------------------------------------------------------------------
# Phase E: Screen reader announcement (model-level validation)
# ---------------------------------------------------------------------------


class TestScreenReaderModel:
    """Tests for the table model fields that feed the TypeScript
    buildAnnouncement() function in TableDetailPanel.tsx.

    These verify the Python-side data is correct (is_row_header, header_level,
    header_col_count) so the frontend has what it needs to build announcements.
    """

    def test_header_cells_have_is_header_true(self):
        table = _make_minimal_table(has_header_row=True)
        header_row = next(r for r in table.rows if r.is_header_row)
        assert all(c.is_header for c in header_row.cells)

    def test_row_header_cells_have_is_row_header_true(self):
        table = _make_minimal_table(has_header_row=True)
        data_row = next(r for r in table.rows if not r.is_header_row)
        data_row.cells[0].is_row_header = True
        table.header_col_count = 1
        assert data_row.cells[0].is_row_header is True
        assert data_row.cells[1].is_row_header is False

    def test_header_level_set_on_primary_header(self):
        from src.tables.table_extractor import _convert_table

        class _FakeCell:
            def __init__(self, text): self.text = text

        class _FakeTable:
            row_count = 2
            col_count = 2
            cells = [None] * 4
            bbox = None

            def extract(self):
                return [["Name", "Value"], ["Alice", "42"]]

        ft = _FakeTable()
        result = _convert_table(ft, page_number=1, index=0, cell_signals={})
        assert result is not None
        header_cells = [c for r in result.rows if r.is_header_row for c in r.cells]
        # All cells in the header row should have header_level >= 1
        assert all(c.header_level >= 1 for c in header_cells)

    def test_confidence_exposed_in_api(self, client):
        table = _make_minimal_table(confidence=0.6)
        job_id = _make_job_with_table(table)
        resp = client.get(f"/api/documents/{job_id}/tables")
        data = resp.json()["tables"][0]
        assert abs(data["confidence"] - 0.6) < 0.01

    def test_header_col_count_exposed_in_api(self, client):
        table = _make_minimal_table()
        table.header_col_count = 1
        job_id = _make_job_with_table(table)
        resp = client.get(f"/api/documents/{job_id}/tables")
        data = resp.json()["tables"][0]
        assert data["header_col_count"] == 1


# ---------------------------------------------------------------------------
# Phase A: Merged cell span detection from PyMuPDF None pattern
# ---------------------------------------------------------------------------


class TestCellSpanDetection:
    """Verify _convert_table() populates col_span / row_span when PyMuPDF
    returns None for cells consumed by a horizontal or vertical merge."""

    def _make_col_span_table(self):
        """2 rows × 3 cols; cell (0,0) spans columns 0 and 1."""
        class _FakeColSpanTable:
            row_count = 2
            col_count = 3
            bbox = (0, 0, 300, 40)
            # flat cells list: (0,0)=anchor, (0,1)=None (col-span), (0,2)=normal
            # row 1: all present
            cells = [
                (0, 0, 200, 20),    # (0,0) anchor — spans cols 0-1
                None,               # (0,1) consumed by (0,0)
                (200, 0, 300, 20),  # (0,2) normal
                (0, 20, 100, 40),   # (1,0) normal
                (100, 20, 200, 40), # (1,1) normal
                (200, 20, 300, 40), # (1,2) normal
            ]
            def extract(self):
                return [["Merged Header", None, "Col C"], ["A", "B", "C"]]
        return _FakeColSpanTable()

    def _make_row_span_table(self):
        """3 rows × 2 cols; cell (0,0) spans rows 0 and 1."""
        class _FakeRowSpanTable:
            row_count = 3
            col_count = 2
            bbox = (0, 0, 200, 60)
            cells = [
                (0, 0, 100, 40),    # (0,0) anchor — spans rows 0-1
                (100, 0, 200, 20),  # (0,1) normal
                None,               # (1,0) consumed by (0,0) row-span
                (100, 20, 200, 40), # (1,1) normal
                (0, 40, 100, 60),   # (2,0) normal
                (100, 40, 200, 60), # (2,1) normal
            ]
            def extract(self):
                return [["Category", "Val1"], [None, "Val2"], ["Other", "Val3"]]
        return _FakeRowSpanTable()

    def test_col_span_detected_from_none_pattern(self):
        from src.tables.table_extractor import _convert_table
        result = _convert_table(self._make_col_span_table(), page_number=1, index=0)
        assert result is not None
        anchor = result.rows[0].cells[0]
        assert anchor.col_span == 2, f"Expected col_span=2, got {anchor.col_span}"
        assert anchor.row_span == 1

    def test_col_span_anchor_text_preserved(self):
        from src.tables.table_extractor import _convert_table
        result = _convert_table(self._make_col_span_table(), page_number=1, index=0)
        assert result is not None
        assert result.rows[0].cells[0].text == "Merged Header"

    def test_col_span_consumed_cell_is_empty(self):
        from src.tables.table_extractor import _convert_table
        result = _convert_table(self._make_col_span_table(), page_number=1, index=0)
        assert result is not None
        # Consumed cell at (0,1) should have empty text and span=1
        consumed = result.rows[0].cells[1]
        assert consumed.text == ""
        assert consumed.col_span == 1

    def test_row_span_detected_from_none_column(self):
        from src.tables.table_extractor import _convert_table
        result = _convert_table(self._make_row_span_table(), page_number=1, index=0)
        assert result is not None
        anchor = result.rows[0].cells[0]
        assert anchor.row_span == 2, f"Expected row_span=2, got {anchor.row_span}"
        assert anchor.col_span == 1

    def test_span_consumed_cells_not_penalised_in_confidence(self):
        from src.tables.table_extractor import _convert_table
        # A table with merges should not have low confidence solely because
        # of the None cells that are part of the merge.
        result = _convert_table(self._make_col_span_table(), page_number=1, index=0)
        assert result is not None
        # Confidence should not drop below 0.7 just because of the merge None.
        assert result.confidence >= 0.7

    def test_no_span_when_all_cells_present(self):
        from src.tables.table_extractor import _convert_table
        class _NoSpanTable:
            row_count = 2
            col_count = 2
            bbox = (0, 0, 200, 40)
            cells = [
                (0, 0, 100, 20), (100, 0, 200, 20),
                (0, 20, 100, 40), (100, 20, 200, 40),
            ]
            def extract(self):
                return [["A", "B"], ["1", "2"]]
        result = _convert_table(_NoSpanTable(), page_number=1, index=0)
        assert result is not None
        for row in result.rows:
            for cell in row.cells:
                assert cell.col_span == 1
                assert cell.row_span == 1


# ---------------------------------------------------------------------------
# Phase C: TABLE_006 — merged cells Markdown fidelity warning
# ---------------------------------------------------------------------------


class TestTable006MergedCells:
    """TABLE_006 fires when a table has cells with col_span or row_span > 1,
    warning that Markdown pipe tables cannot represent the merge."""

    def _issues(self, table: Table) -> list:
        doc = _make_doc(table)
        return validate_document(doc)

    def test_TABLE_006_fires_when_col_span_present(self):
        table = _make_minimal_table(has_caption=True, has_summary=True, has_header_row=True)
        table.rows[0].cells[0].col_span = 2
        issues = self._issues(table)
        rule_ids = [i.rule_id for i in issues]
        assert "TABLE_006" in rule_ids
        match = next(i for i in issues if i.rule_id == "TABLE_006")
        assert match.severity == Severity.WARNING
        assert "merged" in match.message.lower()

    def test_TABLE_006_fires_when_row_span_present(self):
        table = _make_minimal_table(has_caption=True, has_summary=True, has_header_row=True)
        table.rows[0].cells[0].row_span = 2
        issues = self._issues(table)
        assert any(i.rule_id == "TABLE_006" for i in issues)

    def test_TABLE_006_absent_when_no_merges(self):
        table = _make_minimal_table(has_caption=True, has_summary=True, has_header_row=True)
        issues = self._issues(table)
        assert not any(i.rule_id == "TABLE_006" for i in issues)

    def test_TABLE_006_page_number_matches_table(self):
        table = _make_minimal_table(has_caption=True, has_summary=True, has_header_row=True)
        table.page_number = 5
        table.rows[0].cells[0].col_span = 2
        issues = self._issues(table)
        match = next((i for i in issues if i.rule_id == "TABLE_006"), None)
        assert match is not None
        assert match.page_number == 5


# ---------------------------------------------------------------------------
# Phase D: Cell text editing via PATCH /tables/{id}
# ---------------------------------------------------------------------------


class TestCellTextEditAPI:
    """Verify the PATCH endpoint accepts a 'cells' payload and updates
    individual cell text values in the in-memory table model."""

    def test_patch_cells_updates_text(self, client):
        table = _make_minimal_table(has_header_row=True)
        job_id = _make_job_with_table(table)
        resp = client.patch(
            f"/api/documents/{job_id}/tables/{table.table_id}",
            json={"cells": [{"row_index": 0, "col_index": 0, "text": "Updated Header"}]},
        )
        assert resp.status_code == 200
        data = resp.json()
        header_row = next(r for r in data["rows"] if r["is_header_row"])
        assert header_row["cells"][0]["text"] == "Updated Header"

    def test_patch_cells_updates_multiple_cells(self, client):
        table = _make_minimal_table(has_header_row=True)
        job_id = _make_job_with_table(table)
        client.patch(
            f"/api/documents/{job_id}/tables/{table.table_id}",
            json={
                "cells": [
                    {"row_index": 0, "col_index": 0, "text": "Name"},
                    {"row_index": 0, "col_index": 1, "text": "Score"},
                    {"row_index": 1, "col_index": 0, "text": "Alice"},
                    {"row_index": 1, "col_index": 1, "text": "99"},
                ]
            },
        )
        resp = client.get(f"/api/documents/{job_id}/tables")
        rows = resp.json()["tables"][0]["rows"]
        flat = {(c["row_index"], c["col_index"]): c["text"] for r in rows for c in r["cells"]}
        assert flat[(0, 0)] == "Name"
        assert flat[(0, 1)] == "Score"
        assert flat[(1, 0)] == "Alice"
        assert flat[(1, 1)] == "99"

    def test_patch_cells_out_of_range_silently_ignored(self, client):
        table = _make_minimal_table(has_header_row=True)
        job_id = _make_job_with_table(table)
        resp = client.patch(
            f"/api/documents/{job_id}/tables/{table.table_id}",
            json={"cells": [{"row_index": 99, "col_index": 99, "text": "ghost"}]},
        )
        assert resp.status_code == 200
        resp2 = client.get(f"/api/documents/{job_id}/tables")
        for row in resp2.json()["tables"][0]["rows"]:
            for cell in row["cells"]:
                assert cell["text"] != "ghost"

    def test_patch_cells_sets_status_reviewed(self, client):
        table = _make_minimal_table(has_header_row=True)
        assert table.status == TableStatus.AUTO_DETECTED
        job_id = _make_job_with_table(table)
        resp = client.patch(
            f"/api/documents/{job_id}/tables/{table.table_id}",
            json={"cells": [{"row_index": 1, "col_index": 0, "text": "edited"}]},
        )
        assert resp.json()["status"] == "reviewed"

    def test_patch_cells_without_cells_field_preserves_existing(self, client):
        table = _make_minimal_table(has_header_row=True)
        job_id = _make_job_with_table(table)
        # First patch to set text
        client.patch(
            f"/api/documents/{job_id}/tables/{table.table_id}",
            json={"cells": [{"row_index": 0, "col_index": 0, "text": "Set Once"}]},
        )
        # Second patch with no cells field — should not overwrite
        client.patch(
            f"/api/documents/{job_id}/tables/{table.table_id}",
            json={"caption": "New caption"},
        )
        resp = client.get(f"/api/documents/{job_id}/tables")
        rows = resp.json()["tables"][0]["rows"]
        header_cell = next(
            c for r in rows if r["is_header_row"] for c in r["cells"] if c["col_index"] == 0
        )
        assert header_cell["text"] == "Set Once"
